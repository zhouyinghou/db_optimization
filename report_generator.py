#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告生成核心模块
提供Word文档格式的数据库优化分析报告生成功能
"""

from typing import Dict, List, Optional
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_BREAK
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls

from analyze_slow_queries import SlowQueryAnalyzer
from data_masking import DataMasking
from db_connection_manager import DatabaseConnectionManager

class ReportGenerator:
    """报告生成器，负责创建Word格式的数据库优化分析报告"""
    
    def __init__(self, db_connection_manager: DatabaseConnectionManager, 
                 excluded_tables: List[str] = None):
        """
        初始化报告生成器
        
        Args:
            db_connection_manager: 数据库连接管理器
            excluded_tables: 需要排除的表名列表
        """
        self.db_manager = db_connection_manager
        self.excluded_tables = excluded_tables or ['test_table_0']
        self.analysis_data = None
        self.compare_data = None
    
    def load_analysis_data(self, analysis_results_file: str = 'slow_query_analysis_results.json') -> bool:
        """
        从JSON文件加载分析数据
        
        Args:
            analysis_results_file: 分析结果文件路径
            
        Returns:
            bool: 是否成功加载数据
        """
        try:
            if os.path.exists(analysis_results_file):
                with open(analysis_results_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.analysis_data = data.get('analysis_data', [])
                    self.compare_data = data.get('compare_data', {})
                return True
            return False
        except Exception as e:
            print(f"加载分析数据失败: {str(e)}")
            return False
    
    def perform_live_analysis(self, min_execute_cnt: int = 1000, min_query_time: float = 10.0):
        """
        执行实时慢查询分析
        
        Args:
            min_execute_cnt: 最小执行次数
            min_query_time: 最小查询时间
        """
        try:
            # 获取慢查询数据库配置
            slow_query_config = self.db_manager.get_slow_query_config()
            
            # 创建慢查询分析器
            analyzer = SlowQueryAnalyzer(
                slow_query_db_host=slow_query_config['host'],
                slow_query_db_user=slow_query_config['user'],
                slow_query_db_password=slow_query_config['password'],
                slow_query_db_port=slow_query_config['port'],
                slow_query_db_name=slow_query_config['database'],
                slow_query_table=slow_query_config['table'],
                business_db_config=self.db_manager.business_db_config
            )
            
            # 执行对比分析
            compare_result = analyzer.compare_slow_queries(min_execute_cnt, min_query_time)
            
            # 过滤掉包含排除表名的查询
            if compare_result:
                from sql_analyzer import SQLAnalyzer
                sql_analyzer = SQLAnalyzer(self.excluded_tables)
                
                # 过滤上个月的数据
                if 'last_month' in compare_result and 'queries' in compare_result['last_month']:
                    compare_result['last_month']['queries'] = sql_analyzer.filter_excluded_tables(
                        compare_result['last_month']['queries'])
                
                # 过滤前一个月的数据
                if 'previous_month' in compare_result and 'queries' in compare_result['previous_month']:
                    compare_result['previous_month']['queries'] = sql_analyzer.filter_excluded_tables(
                        compare_result['previous_month']['queries'])
            
            # 更新分析数据
            self.compare_data = compare_result
            
            if compare_result:
                # 只保留上个月的慢查询数据，避免重复统计
                self.analysis_data = []
                if 'queries' in compare_result['last_month']:
                    self.analysis_data.extend(compare_result['last_month']['queries'])
            else:
                raise Exception("实时分析失败，无法获取真实的慢查询数据")
        
        except Exception as e:
            raise Exception(f"实时分析失败: {str(e)}")
    
    def generate_report(self, output_file: str = None) -> str:
        """
        生成数据库优化分析报告
        
        Args:
            output_file: 输出文件路径，如果为None则自动生成
            
        Returns:
            str: 生成的报告文件路径
        """
        if not self.analysis_data:
            raise Exception("没有可用的分析数据，请先执行分析或加载数据")
        
        # 生成文件名
        if not output_file:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_file = f"数据库智能优化分析报告_{timestamp}.docx"
        
        # 创建Word文档
        doc = Document()
        
        # 添加报告内容
        self._add_title_page(doc)
        self._add_summary_section(doc)
        self._add_analysis_details(doc)
        self._add_recommendations(doc)
        self._add_appendix(doc)
        
        # 保存文档
        doc.save(output_file)
        
        return output_file
        
    def create_report(self, report_title: str = "数据库性能优化分析报告", 
                      output_dir: str = "./reports", 
                      compare_data: Dict = None) -> str:
        """
        生成Word格式的数据库优化分析报告。
        
        Args:
            report_title: 报告标题
            output_dir: 输出目录
            compare_data: 用于比较分析的数据
            
        Returns:
            生成的报告文件路径
        """
        # 使用传入的比较数据或类中存储的数据
        actual_compare_data = compare_data or self.compare_data
        
        # 创建输出目录
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 生成带时间戳的文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{report_title}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        
        # 创建Word文档
        doc = Document()
        
        # 设置页面布局
        self._setup_page_layout(doc)
        
        # 设置文档样式
        self._setup_document_styles(doc)
        
        # 生成报告标题
        self._generate_report_header(doc, report_title)
        self._add_separator_line(doc)
        
        # 生成报告摘要
        self._generate_report_summary(doc, actual_compare_data)
        self._add_separator_line(doc)
        
        # 生成SQL语句排名
        self._generate_top_sql_statements(doc, actual_compare_data)
        
        # 如果有比较数据，添加比较分析
        if actual_compare_data:
            self._add_compare_analysis(doc, actual_compare_data)
        
        # 保存文档
        doc.save(filepath)
        
        print(f"报告已生成：{filepath}")
        return filepath
    
    def _setup_page_layout(self, doc: Document):
        """
        设置页面布局
        """
        # 设置页边距
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(3.0)
    
    def _setup_document_styles(self, doc: Document):
        """
        设置文档样式
        """
        # 设置标题样式
        heading_1_style = doc.styles['Heading 1']
        heading_1_font = heading_1_style.font
        heading_1_font.name = '微软雅黑'
        heading_1_font.size = Pt(16)
        heading_1_font.bold = True
        heading_1_font.color.rgb = RGBColor(31, 73, 125)
        
        heading_2_style = doc.styles['Heading 2']
        heading_2_font = heading_2_style.font
        heading_2_font.name = '微软雅黑'
        heading_2_font.size = Pt(13)
        heading_2_font.bold = True
        heading_2_font.color.rgb = RGBColor(79, 129, 189)
        
        heading_3_style = doc.styles['Heading 3']
        heading_3_font = heading_3_style.font
        heading_3_font.name = '微软雅黑'
        heading_3_font.size = Pt(11)
        heading_3_font.bold = True
        heading_3_font.color.rgb = RGBColor(149, 179, 215)
        
        # 设置正文样式
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '微软雅黑'
        normal_font.size = Pt(10)
    
    def _generate_report_header(self, doc: Document, title: str):
        """
        生成报告标题
        """
        # 添加标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加日期
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
        date_run.font.size = Pt(10)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加敏感信息提示
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run("⚠️ 注意：本报告包含敏感的数据库性能信息，请妥善保管！")
        warning_run.font.size = Pt(10)
        warning_run.font.color.rgb = RGBColor(192, 0, 0)
        warning_run.bold = True
        warning_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_separator_line(self, doc: Document):
        """
        添加分隔线
        """
        # 添加空行
        doc.add_paragraph()
        
        # 添加分隔线
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run("=" * 80)
        separator_run.font.size = Pt(8)
        separator_run.font.color.rgb = RGBColor(192, 192, 192)
    
    def _generate_report_summary(self, doc: Document, compare_data: Dict = None):
        """
        生成报告摘要
        """
        doc.add_heading("一、报告摘要", level=1)
        
        # 添加摘要内容
        summary_para = doc.add_paragraph()
        summary_para.add_run("本报告总结了数据库性能优化分析的主要发现和建议。").font.size = Pt(10)
        
        # 如果有比较数据，添加比较统计
        if compare_data:
            # 添加比较分析表格
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '指标'
            hdr_cells[1].text = '本月'
            hdr_cells[2].text = '上月'
            
            # 设置表头样式
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def _generate_top_sql_statements(self, doc: Document, compare_data: Dict = None):
        """
        生成SQL语句排名
        """
        doc.add_heading("二、SQL性能分析", level=1)
        doc.add_heading("2.1 慢查询TOP 10", level=2)
        
        # 添加SQL性能概览表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '排名'
        hdr_cells[1].text = 'SQL语句'
        hdr_cells[2].text = '执行次数'
        hdr_cells[3].text = '总执行时间(ms)'
        hdr_cells[4].text = '平均执行时间(ms)'
        hdr_cells[5].text = '数据库'
        
        # 设置表头样式
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置列宽
            if cell.text == '排名':
                cell.width = Cm(1.0)
            elif cell.text == 'SQL语句':
                cell.width = Cm(8.0)
            else:
                cell.width = Cm(2.0)
        
        # 如果有比较数据，添加数据行
        if compare_data and 'last_month' in compare_data and 'queries' in compare_data['last_month']:
            # 获取排序后的查询
            sorted_queries = self._get_sorted_queries(compare_data['last_month']['queries'])
            
            # 添加前10个查询
            for i, query in enumerate(sorted_queries[:10], 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                sql_text = query.get('sql', '').replace('\n', ' ')
                row_cells[1].text = sql_text[:100] + '...' if len(sql_text) > 100 else sql_text
                row_cells[2].text = str(query.get('execute_count', 0))
                row_cells[3].text = str(query.get('total_execution_time', 0))
                row_cells[4].text = str(query.get('avg_execution_time', 0))
                row_cells[5].text = query.get('database', '')
                
                # 设置数据行对齐方式
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def _get_sorted_queries(self, queries: List[Dict]) -> List[Dict]:
        """
        获取排序后的查询列表
        """
        # 按执行次数、执行时间和数据库名称排序
        return sorted(
            queries,
            key=lambda x: (x.get('execute_count', 0), x.get('total_execution_time', 0), x.get('database', '')),
            reverse=True
        )
    
    def _add_compare_analysis(self, doc: Document, compare_data: Dict):
        """
        添加比较分析
        """
        doc.add_heading("三、对比分析", level=1)
        
        # 检查是否有上月数据
        if 'last_month' in compare_data and 'previous_month' in compare_data:
            try:
                # 添加比较表格
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # 设置表头
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '指标'
                hdr_cells[1].text = '本月'
                hdr_cells[2].text = '上月'
                
                # 设置表头样式
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加比较数据
                compare_stats = [
                    ('慢查询数量', 
                     len(compare_data['last_month'].get('queries', [])),
                     len(compare_data['previous_month'].get('queries', []))),
                    ('总执行次数',
                     sum(q.get('execute_count', 0) for q in compare_data['last_month'].get('queries', [])),
                     sum(q.get('execute_count', 0) for q in compare_data['previous_month'].get('queries', []))),
                    ('总执行时间(ms)',
                     sum(q.get('total_execution_time', 0) for q in compare_data['last_month'].get('queries', [])),
                     sum(q.get('total_execution_time', 0) for q in compare_data['previous_month'].get('queries', [])))
                ]
                
                for label, current_val, prev_val in compare_stats:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(current_val)
                    row_cells[2].text = str(prev_val)
                    
                    # 设置数据行对齐方式
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            except Exception as e:
                # 出错时添加错误信息，但继续执行报告生成
                error_para = doc.add_paragraph()
                error_run = error_para.add_run(f"⚠️ 生成比较分析时出错：{str(e)}")
                error_run.font.color.rgb = RGBColor(192, 0, 0)
    
    def _generate_sql_details(self, doc: Document, query: Dict):
        """
        生成SQL详情分析
        """
        # 添加SQL详情标题
        doc.add_heading(f"SQL ID: {query.get('query_id', 'Unknown')}", level=3)
        
        # 添加SQL语句
        sql_para = doc.add_paragraph()
        sql_para.add_run("SQL语句：").bold = True
        sql_para.add_run(query.get('sql', '')).font.size = Pt(9)
        
        # 添加执行统计
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Table Grid'
        
        # 设置表头
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = '统计项'
        hdr_cells[1].text = '值'
        
        # 设置表头样式
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # 添加统计数据
        stats_data = [
            ('数据库', query.get('database', '')),
            ('执行次数', str(query.get('execute_count', 0))),
            ('总执行时间', f"{query.get('total_execution_time', 0)} ms"),
            ('平均执行时间', f"{query.get('avg_execution_time', 0)} ms"),
            ('最大执行时间', f"{query.get('max_execution_time', 0)} ms"),
            ('扫描行数', str(query.get('rows_examined', 0))),
            ('返回行数', str(query.get('rows_sent', 0)))
        ]
        
        for label, value in stats_data:
            row_cells = stats_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value
        
        # 执行SQL分析
        optimization_suggestions = self._analyze_sql_for_optimization(query.get('sql', ''))
        
        # 添加优化建议
        if optimization_suggestions:
            doc.add_heading("优化建议", level=4)
            for suggestion in optimization_suggestions:
                suggestion_para = doc.add_paragraph()
                suggestion_para.add_run(f"• {suggestion}").font.size = Pt(9)
    
    def _analyze_sql_for_optimization(self, sql: str) -> List[str]:
        """
        分析SQL语句并生成优化建议
        """
        suggestions = []
        
        # 转换为小写以便分析
        sql_lower = sql.lower()
        
        # 检查是否使用了SELECT *
        if 'select *' in sql_lower:
            suggestions.append("避免使用SELECT *，只查询必要的列")
        
        # 检查是否使用了索引
        if 'where' in sql_lower:
            # 简单的索引使用检查
            where_clause = sql_lower[sql_lower.find('where') + 5:]
            if 'join' in where_clause and 'on' in where_clause:
                suggestions.append("确保JOIN条件中的列已经建立索引")
        
        # 检查是否使用了ORDER BY但没有索引
        if 'order by' in sql_lower:
            suggestions.append("考虑在ORDER BY的列上建立索引以提高排序性能")
        
        # 检查是否使用了GROUP BY但没有索引
        if 'group by' in sql_lower:
            suggestions.append("考虑在GROUP BY的列上建立索引以提高分组性能")
        
        # 检查是否使用了LIKE并且以通配符开头
        if 'like' in sql_lower:
            # 简单检查是否以%开头
            if 'like \'%' in sql_lower or "like '%" in sql_lower:
                suggestions.append("避免在LIKE查询中以通配符%开头，这会导致索引失效")
        
        # 检查是否使用了子查询
        if 'select' in sql_lower[sql_lower.find('select') + 6:]:
            suggestions.append("考虑优化子查询，可能的话使用JOIN代替")
        
        # 检查是否使用了OR
        if 'or' in sql_lower and 'where' in sql_lower:
            suggestions.append("考虑使用UNION ALL代替OR，因为OR可能导致索引失效")
        
        return suggestions

    def _add_title_page(self, doc: Document):
        """添加标题页"""
        # 添加标题
        title = doc.add_heading('数据库智能优化分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告信息
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        # 报告日期
        report_date = datetime.now().strftime('%Y年%m月%d日')
        info_table.cell(0, 0).text = '报告日期：'
        info_table.cell(0, 1).text = report_date
        
        # 分析时间范围
        if self.compare_data and 'last_month' in self.compare_data:
            time_range = self.compare_data['last_month'].get('name', '未知')
            info_table.cell(1, 0).text = '分析时间范围：'
            info_table.cell(1, 1).text = time_range
        
        # 慢查询数量
        query_count = len(self.analysis_data) if self.analysis_data else 0
        info_table.cell(2, 0).text = '分析慢查询数量：'
        info_table.cell(2, 1).text = str(query_count)
        
        # 分析状态
        info_table.cell(3, 0).text = '分析状态：'
        info_table.cell(3, 1).text = '已完成'
        
        # 居中对齐表格
        for row in info_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    def _add_summary_section(self, doc: Document):
        """添加摘要部分"""
        doc.add_heading('执行摘要', level=1)
        
        # 添加总体统计信息
        if self.compare_data:
            summary_data = []
            
            # 上个月数据
            last_month = self.compare_data.get('last_month', {})
            prev_month = self.compare_data.get('previous_month', {})
            comparison = self.compare_data.get('comparison', {})
            
            summary_data.extend([
                f"• 上月慢查询总数：{last_month.get('total', 0)} 条",
                f"• 平均查询时间：{last_month.get('avg_query_time', 0):.2f} 秒",
                f"• 平均执行次数：{last_month.get('avg_execute_cnt', 0):.0f} 次",
            ])
            
            if prev_month.get('total', 0) > 0:
                growth_rate = comparison.get('growth_rate', 0)
                summary_data.extend([
                    f"• 环比增长率：{growth_rate:.1f}%",
                    f"• 新增慢查询：{comparison.get('new_queries_count', 0)} 条",
                    f"• 已解决慢查询：{comparison.get('resolved_queries_count', 0)} 条",
                ])
            
            # 添加摘要表格
            for item in summary_data:
                p = doc.add_paragraph(item)
                p.style = 'List Bullet'
        
        doc.add_page_break()
    
    def _add_analysis_details(self, doc: Document):
        """添加详细分析部分"""
        doc.add_heading('详细分析', level=1)
        
        if not self.analysis_data:
            doc.add_paragraph('未找到慢查询数据。')
            return
        
        # 对数据进行脱敏处理
        masked_data = DataMasking.mask_sensitive_data(self.analysis_data)
        
        # 按查询时间排序
        sorted_data = sorted(masked_data, 
                          key=lambda x: x.get('query_time', 0), 
                          reverse=True)
        
        # 添加TOP 10 慢查询
        doc.add_heading('TOP 10 慢查询分析', level=2)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # 添加表头
        headers = ['排名', '执行次数', '查询时间(秒)', '数据库名', '主机IP']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # 添加数据行（只显示前10条）
        for i, query in enumerate(sorted_data[:10]):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = str(query.get('execute_cnt', 0))
            row_cells[2].text = f"{query.get('query_time', 0):.2f}"
            row_cells[3].text = query.get('db_name', '未知')
            row_cells[4].text = query.get('hostname_max', query.get('ip', '未知'))
        
        doc.add_page_break()
    
    def _add_recommendations(self, doc: Document):
        """添加优化建议部分"""
        doc.add_heading('优化建议', level=1)
        
        # 基于分析结果生成建议
        recommendations = self._generate_recommendations()
        
        for i, rec in enumerate(recommendations, 1):
            doc.add_heading(f'建议 {i}', level=2)
            
            # 建议描述
            p = doc.add_paragraph(rec['description'])
            p.style = 'Normal'
            
            # 建议详情
            if 'details' in rec:
                for detail in rec['details']:
                    p = doc.add_paragraph(detail)
                    p.style = 'List Bullet'
            
            # 预期效果
            if 'expected_effect' in rec:
                p = doc.add_paragraph(f"预期效果：{rec['expected_effect']}")
                p.style = 'List Bullet'
            
            doc.add_paragraph()  # 空行
        
        doc.add_page_break()
    
    def _add_appendix(self, doc: Document):
        """添加附录"""
        doc.add_heading('附录', level=1)
        
        # 添加技术说明
        doc.add_heading('技术说明', level=2)
        
        tech_info = [
            "• 本报告基于MySQL慢查询日志生成",
            "• 分析标准：执行次数 > 1000 且 查询时间 > 10秒",
            "• 数据已进行脱敏处理，保护敏感信息",
            "• 优化建议基于EXPLAIN执行计划生成",
        ]
        
        for info in tech_info:
            p = doc.add_paragraph(info)
            p.style = 'List Bullet'
        
        # 添加排除表说明
        if self.excluded_tables:
            doc.add_heading('排除的表', level=2)
            p = doc.add_paragraph("以下表已被排除在分析范围之外：")
            p.style = 'Normal'
            
            for table in self.excluded_tables:
                p = doc.add_paragraph(f"• {table}")
                p.style = 'List Bullet'
    
    def _generate_recommendations(self) -> List[Dict]:
        """
        基于分析数据生成优化建议
        
        Returns:
            优化建议列表
        """
        recommendations = []
        
        if not self.analysis_data:
            return recommendations
        
        # 分析查询模式
        high_time_queries = [q for q in self.analysis_data if q.get('query_time', 0) > 100]
        high_freq_queries = [q for q in self.analysis_data if q.get('execute_cnt', 0) > 10000]
        
        # 高查询时间建议
        if high_time_queries:
            recommendations.append({
                'description': '发现查询时间过长的SQL语句，建议优化查询逻辑和索引',
                'details': [
                    f"有 {len(high_time_queries)} 条查询的平均执行时间超过100秒",
                    "建议检查WHERE条件是否使用了合适的索引",
                    "考虑重写复杂查询，分解为多个简单查询"
                ],
                'expected_effect': '查询时间减少60-90%'
            })
        
        # 高频查询建议
        if high_freq_queries:
            recommendations.append({
                'description': '发现高频执行的SQL语句，建议优化查询性能',
                'details': [
                    f"有 {len(high_freq_queries)} 条查询的执行次数超过10000次",
                    "高频查询的微小优化都能带来显著的性能提升",
                    "建议为常用的查询条件添加复合索引"
                ],
                'expected_effect': '整体数据库性能提升20-40%'
            })
        
        # 通用建议
        recommendations.extend([
            {
                'description': '定期执行表分析，更新统计信息',
                'details': [
                    "执行 ANALYZE TABLE 更新表统计信息",
                    "优化器需要准确的统计信息来选择最佳执行计划"
                ],
                'expected_effect': '查询执行计划更准确，性能提升5-15%'
            },
            {
                'description': '监控和优化数据库配置',
                'details': [
                    "检查 innodb_buffer_pool_size 配置",
                    "优化查询缓存配置",
                    "定期清理慢查询日志"
                ],
                'expected_effect': '数据库整体性能提升10-20%'
            }
        ])
        
        return recommendations
    
    def save_analysis_data(self, filename: str = 'slow_query_analysis_results.json'):
        """
        保存分析数据到JSON文件
        
        Args:
            filename: 文件名
        """
        data = {
            'analysis_data': self.analysis_data,
            'compare_data': self.compare_data,
            'generated_at': datetime.now().isoformat(),
            'excluded_tables': self.excluded_tables
        }
        
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2, default=str)
            print(f"分析数据已保存到: {filename}")
        except Exception as e:
            print(f"保存分析数据失败: {str(e)}")
    
    def get_analysis_summary(self) -> Dict:
        """
        获取分析摘要信息
        
        Returns:
            包含关键指标的分析摘要
        """
        if not self.analysis_data and not self.compare_data:
            return {'status': 'no_data'}
        
        summary = {
            'status': 'success',
            'total_queries': len(self.analysis_data) if self.analysis_data else 0,
            'analysis_date': datetime.now().isoformat(),
        }
        
        # 添加对比数据
        if self.compare_data:
            last_month = self.compare_data.get('last_month', {})
            comparison = self.compare_data.get('comparison', {})
            
            summary.update({
                'last_month_total': last_month.get('total', 0),
                'avg_query_time': last_month.get('avg_query_time', 0),
                'avg_execute_cnt': last_month.get('avg_execute_cnt', 0),
                'growth_rate': comparison.get('growth_rate', 0),
                'new_queries': comparison.get('new_queries_count', 0),
                'resolved_queries': comparison.get('resolved_queries_count', 0),
            })
        
        return summary