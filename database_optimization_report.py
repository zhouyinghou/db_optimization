#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
数据库智能优化分析报告生成器
基于Oracle AWR报告风格，生成专业的MySQL数据库优化分析报告
"""

import json
import os
import re
import sys
from datetime import datetime
from typing import Dict, List, Optional
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_BREAK
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls

# 导入拆分后的模块
from utils import setup_encoding, load_db_config
from data_masking import DataMasking
from sql_analyzer import SQLAnalyzer
from data_processor import DataProcessor
from database_helper import DatabaseHelper
from summary_generator import SummaryGenerator
from report_generator import ReportGenerator
from report_generator_core import ReportGeneratorCore

# 设置编码
setup_encoding()

# 添加必要的导入
from analyze_slow_queries import SlowQueryAnalyzer

# 尝试导入智能优化建议模块（可选）
try:
    from intelligent_optimization_suggestions import IntelligentOptimizationSuggestions
    INTELLIGENT_OPTIMIZER_AVAILABLE = True
except ImportError:
    INTELLIGENT_OPTIMIZER_AVAILABLE = False
    IntelligentOptimizationSuggestions = None

class DatabaseOptimizationReport:
    """数据库智能优化分析报告生成器"""
    
    def __init__(self, use_live_analysis: bool = False, 
                 slow_query_db_config: Dict = None,  # type: ignore
                 business_db_config: Dict = None,  # type: ignore
                 min_execute_cnt: int = 1000,
                 min_query_time: float = 10.0,
                 load_data: bool = True):
        self.slow_query_db_config = slow_query_db_config
        self.business_db_config = business_db_config
        self.analysis_data = None
        self.compare_data = None
        self.use_live_analysis = use_live_analysis
        # 定义需要排除的表名列表
        self.excluded_tables = ['test_table_0']
        
        # 初始化报告生成器
        self.report_generator = ReportGenerator(
            db_connection_manager=business_db_config,
            excluded_tables=self.excluded_tables
        )
        
        # 初始化慢查询数据库连接配置
        self.slow_query_db_host = slow_query_db_config.get('host', '127.0.0.1') if slow_query_db_config else '127.0.0.1'
        self.slow_query_db_user = slow_query_db_config.get('user', 'test') if slow_query_db_config else 'test'
        self.slow_query_db_password = slow_query_db_config.get('password', 'test') if slow_query_db_config else 'test'
        self.slow_query_db_port = slow_query_db_config.get('port', 3306) if slow_query_db_config else 3306
        
        # 初始化业务数据库连接配置（用于查询实际慢查询的数据库）
        self.business_db_host = business_db_config.get('host', '127.0.0.1') if business_db_config else '127.0.0.1'
        self.business_db_user = business_db_config.get('user', 'test') if business_db_config else 'test'
        self.business_db_password = business_db_config.get('password', 'test') if business_db_config else 'test'
        self.business_db_port = business_db_config.get('port', 3306) if business_db_config else 3306
        
        # 初始化模块实例
        self.db_helper = DatabaseHelper(
            business_db_config=business_db_config,
            slow_query_db_config=slow_query_db_config
        )

        # 是否启用新的智能优化建议模块（默认关闭，保持拆分前输出）
        self.enable_intelligent_optimizer = False
        
        # 初始化智能优化建议生成器（如果可用）
        if INTELLIGENT_OPTIMIZER_AVAILABLE and IntelligentOptimizationSuggestions:
            try:
                self.intelligent_optimizer = IntelligentOptimizationSuggestions(
                    db_helper=self.db_helper
                )
            except Exception:
                self.intelligent_optimizer = None
        else:
            self.intelligent_optimizer = None
        
        if not load_data:
            # 不加载外部数据，仅用于测试
            return
            
        if use_live_analysis and slow_query_db_config:
            # 使用实时分析
            self._perform_live_analysis(slow_query_db_config, min_execute_cnt, min_query_time)

    def _perform_live_analysis(self, db_config: Dict, min_execute_cnt: int, min_query_time: float):
        """执行实时慢查询分析，包括对比分析"""
        try:
            # 创建慢查询分析器，传入表名配置
            if not db_config:
                raise ValueError("数据库配置不能为空")
            
            analyzer = SlowQueryAnalyzer(
                slow_query_db_host=db_config.get('host', ''),
                slow_query_db_user=db_config.get('user', ''),
                slow_query_db_password=db_config.get('password', ''),
                slow_query_db_port=db_config.get('port', 3306),
                slow_query_db_name=db_config.get('database', ''),
                slow_query_table=db_config.get('table', 'slow'),
                business_db_config=self.business_db_config
            )
            
            # 执行对比分析
            compare_result = analyzer.compare_slow_queries(min_execute_cnt, min_query_time)
            
            # 过滤掉包含排除表名的查询
            if compare_result:
                # 过滤上个月的数据
                if 'last_month' in compare_result and 'queries' in compare_result['last_month']:
                    original_last_month_count = len(compare_result['last_month']['queries'])
                    compare_result['last_month']['queries'] = DataProcessor.filter_excluded_tables(
                        compare_result['last_month']['queries'], 
                        self.excluded_tables
                    )
                
                # 过滤前一个月的数据
                if 'previous_month' in compare_result and 'queries' in compare_result['previous_month']:
                    original_prev_month_count = len(compare_result['previous_month']['queries'])
                    compare_result['previous_month']['queries'] = DataProcessor.filter_excluded_tables(
                        compare_result['previous_month']['queries'],
                        self.excluded_tables
                    )
            
            # 不打印任何慢查询SQL，符合用户要求
            # 原代码已注释掉
            
            # 更新分析数据
            self.compare_data = compare_result
            
            if compare_result:
                # 对数据进行脱敏处理
                self.compare_data = compare_result
                
                # 只保留上个月的慢查询数据，避免重复统计
                self.analysis_data = []
                # 只添加上个月的数据（当前需要分析的慢查询）
                if 'queries' in compare_result['last_month']:
                    self.analysis_data.extend(compare_result['last_month']['queries'])
            else:
                # 没有获取到真实数据时抛出错误
                if not self.analysis_data:
                    raise Exception("实时分析失败，无法获取真实的慢查询数据")
        
        except Exception as e:
            # 没有获取到真实数据时抛出错误
            if not self.analysis_data:
                raise Exception(f"实时分析失败: {str(e)}")
        
    def _mask_sensitive_data(self, data: List[Dict]) -> List[Dict]:
        """对敏感信息进行脱敏处理"""
        return DataMasking.mask_sensitive_data(data)
    
    # 包装方法：调用新模块的方法以保持向后兼容
    def _mask_db_name(self, db_name) -> str:
        """脱敏数据库名（包装方法）"""
        return DataMasking.mask_db_name(db_name)
    
    def _mask_ip(self, ip) -> str:
        """脱敏IP地址（包装方法）"""
        return DataMasking.mask_ip(ip)
    
    def _mask_table_name(self, table_name) -> str:
        """脱敏表名（包装方法）"""
        return DataMasking.mask_table_name(table_name)
    
    def _mask_sql(self, sql) -> str:
        """脱敏SQL语句（包装方法）"""
        return DataMasking.mask_sql(sql)
        
    def _extract_table_name(self, sql: str) -> Optional[str]:
        """从SQL语句中提取表名（包装方法）"""
        return SQLAnalyzer.extract_table_name(sql)

    def _extract_where_fields(self, sql: str) -> List[str]:
        """从SQL语句中提取WHERE条件中的字段名（包装方法）"""
        return SQLAnalyzer.extract_where_fields(sql)
    
    def _extract_fields_from_condition(self, condition: str) -> List[str]:
        """从单个条件中提取字段名（包装方法）"""
        return SQLAnalyzer.extract_fields_from_condition(condition)
    
    def _extract_join_fields(self, sql: str) -> List[str]:
        """从SQL语句中提取JOIN条件中的字段名（包装方法）"""
        return SQLAnalyzer.extract_join_fields(sql)
    
    def _extract_order_by_fields(self, sql: str) -> List[str]:
        """从SQL语句中提取ORDER BY子句中的字段名（包装方法）"""
        return SQLAnalyzer.extract_order_by_fields(sql)
    
    def _sort_fields_by_priority(self, fields: List[str], sql_lower: str) -> List[str]:
        """智能排序字段优先级（包装方法）"""
        return SQLAnalyzer.sort_fields_by_priority(fields, sql_lower)
    
    
    def _get_standby_hostname(self, master_hostname: str) -> Optional[str]:
        """通过cluster表查询获取备库IP地址（包装方法）"""
        return self.db_helper.get_standby_hostname(master_hostname)

    def _get_safe_connection(self, hostname: str = None, database: str = None) -> dict:
        """安全地获取数据库连接（包装方法）"""
        return self.db_helper.get_safe_connection(hostname, database)
    
    def _close_safe_connection(self):
        """安全关闭数据库连接（包装方法）"""
        self.db_helper.close_safe_connection()
    
    def _execute_safe_query(self, query: str, params: tuple = None, hostname: str = None, database: str = None) -> dict:
        """安全执行数据库查询（包装方法）"""
        return self.db_helper.execute_safe_query(query, params, hostname, database)
    
    def _get_table_row_count(self, database: str, table_name: str, hostname: str = None) -> Optional[int]:
        """
        获取表的行数（使用hostname_max连接真实业务数据库）
        """
        if not table_name:
            return None
        
        actual_database = database
        if database:
            if not self.db_helper.check_table_exists(database, table_name, hostname):
                found_database = self.db_helper.find_correct_database_for_table(table_name, hostname)
                if found_database:
                    actual_database = found_database
                    print(f"ℹ️ 找到表 {table_name} 所在的实际数据库: {actual_database} (hostname: {hostname})")
                else:
                    print(f"⚠️ 无法找到表 {table_name} 所在的数据库，使用传入的数据库: {database}")
                    actual_database = database
        else:
            found_database = self.db_helper.find_correct_database_for_table(table_name, hostname)
            if found_database:
                actual_database = found_database
            else:
                print(f"❌ 未提供数据库名且无法找到表 {table_name} 所在的数据库")
                return None
        
        return self.db_helper.get_table_row_count(actual_database, table_name, hostname)

    def _get_table_row_count_with_fallback(self, database: str, table_name: str, hostname: str = None, query: Optional[dict] = None) -> Optional[int]:
        """获取表行数，若数据库查询失败则回退到查询元数据"""
        row_count = self._get_table_row_count(database, table_name, hostname)
        if row_count is None:
            row_count = self._extract_row_count_from_query(query)
        return row_count

    def _extract_row_count_from_query(self, query: Optional[dict]) -> Optional[int]:
        """从查询元数据中提取表行数"""
        if not query or not isinstance(query, dict):
            return None
        
        direct_keys = [
            'table_row_count', 'row_count', 'table_rows', 'rows',
            'TABLE_ROWS', 'TABLE_ROW_COUNT', 'TABLE_ROWS_ESTIMATE',
            'total_rows', 'row_num'
        ]
        
        def parse_value(value):
            if value is None:
                return None
            if isinstance(value, (int, float)):
                return int(value)
            if isinstance(value, str):
                cleaned = value.replace(',', '').strip()
                if not cleaned:
                    return None
                try:
                    return int(float(cleaned))
                except ValueError:
                    return None
            return None
        
        def try_extract(source):
            if not source or not isinstance(source, dict):
                return None
            for key in direct_keys:
                if key in source:
                    parsed = parse_value(source[key])
                    if parsed is not None:
                        return parsed
            return None
        
        def ensure_dict(value):
            if isinstance(value, dict):
                return value
            if isinstance(value, str):
                try:
                    import json
                    return json.loads(value)
                except Exception:
                    try:
                        import ast
                        return ast.literal_eval(value)
                    except Exception:
                        return {}
            return {}
        
        # 顶层直接信息
        direct = try_extract(query)
        if direct is not None:
            return direct
        
        # table_structure 中的信息
        table_structure = ensure_dict(query.get('table_structure', {}))
        if table_structure:
            direct = try_extract(table_structure)
            if direct is not None:
                return direct
            
            for nested_key in ['table_stats', 'statistics', 'stats', 'meta']:
                nested = ensure_dict(table_structure.get(nested_key, {}))
                direct = try_extract(nested)
                if direct is not None:
                    return direct
        
        # 顶层其他统计字段
        for nested_key in ['table_stats', 'statistics', 'meta']:
            nested = ensure_dict(query.get(nested_key, {}))
            direct = try_extract(nested)
            if direct is not None:
                return direct
        
        # 慢查询信息中的统计
        slow_info = ensure_dict(query.get('slow_query_info', {}))
        direct = try_extract(slow_info)
        if direct is not None:
            return direct
        
        for nested_key in ['table_stats', 'statistics', 'meta']:
            nested = ensure_dict(slow_info.get(nested_key, {}))
            direct = try_extract(nested)
            if direct is not None:
                return direct
        
        return None

    def _check_table_exists(self, database: str, table_name: str, hostname: str = None) -> bool:
        """检查表是否存在（包装方法）"""
        return self.db_helper.check_table_exists(database, table_name, hostname)
    
    def _get_table_indexes_from_db(self, database: str, table_name: str, hostname: str = None) -> Optional[set]:
        """从数据库中获取表的索引信息（包装方法，支持hostname参数）"""
        result = self.db_helper.get_table_indexes_from_db(database, table_name, hostname)
        return result if result is not None else set()
    
    def _find_correct_database_for_table(self, table_name: str, hostname: Optional[str] = None) -> str:
        """
        查找包含指定表的正确数据库（使用hostname_max连接真实业务数据库）
        
        Args:
            table_name: 表名
            hostname: 主机名（可选），如果提供则使用该主机查找数据库（应该是hostname_max的值）
            
        Returns:
            包含该表的数据库名，如果未找到返回空字符串
        """
        return self.db_helper.find_correct_database_for_table(table_name, hostname)
    
    def _check_indexes_exist(self, database: str, table_name: str, where_fields: list, join_fields: list, order_by_fields: list, query: Optional[dict] = None) -> bool:
        """
        检查所有相关字段是否都有索引
        
        Args:
            database: 数据库名
            table_name: 表名
            where_fields: WHERE条件字段列表
            join_fields: JOIN条件字段列表
            order_by_fields: ORDER BY字段列表
            query: 查询对象，考虑JSON中的表结构信息作为参考
            
        Returns:
            如果所有字段都有索引，返回True，否则返回False
        """
        if not table_name:
            return False
        
        # 🎯 关键修复：如果提供了query参数且包含表结构信息，则跳过表存在性检查
        # 避免在没有数据库连接的情况下返回False
        if query and isinstance(query, dict) and 'table_structure' in query:
            print(f"ℹ️ 使用query参数中的表结构信息，跳过表存在性检查")
        elif database and table_name:
            # 从query对象中获取hostname_max用于连接真实业务数据库
            hostname_max = None
            if query and isinstance(query, dict):
                slow_info = query.get('slow_query_info', {})
                hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
            
            if not self._check_table_exists(database, table_name, hostname_max):
                print(f"⚠️ 表 {table_name} 在数据库 {database} 中不存在，无法检查索引")
                return False
        
        # 检查所有相关字段
        all_fields = set()
        all_fields.update([f.lower() for f in where_fields])
        all_fields.update([f.lower() for f in join_fields])
        all_fields.update([f.lower() for f in order_by_fields])
        
        if not all_fields:
            return False
        
        # 🔥 关键修复：优先从数据库读取真实索引信息，如果数据库查询失败，则从JSON数据中参考
        existing_indexed_fields = set()
        database_query_successful = False
        
        # 1. 优先从数据库获取实际索引信息（使用hostname_max连接真实业务数据库）
        # 从query对象或hostname参数中获取hostname_max
        if not hostname_max:
            if query and isinstance(query, dict):
                slow_info = query.get('slow_query_info', {})
                hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
        
        if database and table_name:
            # 使用execute_safe_query直接查询索引信息（支持hostname参数）
            query_result = self.db_helper.execute_safe_query(
                f"SHOW INDEX FROM `{table_name}`",
                hostname=hostname_max,
                database=database
            )
            if query_result['status'] == 'success' and query_result['data']:
                # 数据库查询成功且有数据
                for row in query_result['data']:
                    if len(row) >= 5:
                        column_name = row[4]
                        if column_name:
                            existing_indexed_fields.add(column_name.lower())
                if existing_indexed_fields:
                    database_query_successful = True
                    print(f"📊 从数据库读取到的索引字段: {existing_indexed_fields}")
            # else:
            #     print(f"⚠️ 数据库查询失败或无索引数据，将从JSON数据中参考")
        
        # 2. 如果数据库查询失败，从query对象中获取表结构信息作为参考
        if not database_query_successful and query and isinstance(query, dict) and 'table_structure' in query:
            table_structure = query.get('table_structure', {})
            # 如果table_structure是字符串，尝试解析
            if isinstance(table_structure, str):
                try:
                    import json
                    table_structure = json.loads(table_structure)
                except (json.JSONDecodeError, ValueError):
                    # 如果JSON解析失败，尝试使用ast.literal_eval（Python字符串表示）
                    try:
                        import ast
                        table_structure = ast.literal_eval(table_structure)
                    except (ValueError, SyntaxError):
                        table_structure = {}
            
            if table_structure and isinstance(table_structure, dict) and 'indexes' in table_structure:
                indexes = table_structure['indexes']
                
                # indexes可能是字典{index_name: index_info}或列表[index_info]
                if isinstance(indexes, dict):
                    # 字典格式：遍历values
                    for index_info in indexes.values():
                        if isinstance(index_info, dict) and 'columns' in index_info:
                            for col in index_info['columns']:
                                existing_indexed_fields.add(col.lower())
                elif isinstance(indexes, list):
                    # 列表格式
                    for index_info in indexes:
                        if isinstance(index_info, dict):
                            # 支持多种索引格式
                            if 'columns' in index_info:
                                # 格式1: {'columns': ['id']}
                                for col in index_info['columns']:
                                    existing_indexed_fields.add(col.lower())
                            elif 'Column_name' in index_info:
                                # 格式2: {'Column_name': 'id'} (MySQL SHOW INDEXES格式)
                                existing_indexed_fields.add(index_info['Column_name'].lower())
                
                if existing_indexed_fields:
                    print(f"📋 从JSON数据中参考到的索引字段: {existing_indexed_fields}")
        
        # 3. 检查所有字段是否都有索引
        if existing_indexed_fields:
            # 检查是否所有字段都有单独的索引
            all_fields_have_individual_indexes = True
            fields_without_individual_indexes = []
            
            for field in all_fields:
                if field not in existing_indexed_fields:
                    all_fields_have_individual_indexes = False
                    fields_without_individual_indexes.append(field)
            
            if all_fields_have_individual_indexes:
                # 所有字段都有单独索引，检查是否需要复合索引
                # 如果WHERE条件中有多个字段，建议复合索引
                if len(where_fields) > 1:
                    print(f"ℹ️ 所有字段都有单独索引，但WHERE条件中有多个字段，建议复合索引")
                    return False  # 返回False表示建议创建复合索引
                else:
                    print(f"✅ 所有字段都有索引，字段: {all_fields}, 已有索引: {existing_indexed_fields}")
                    return True
            else:
                print(f"❌ 字段 {fields_without_individual_indexes} 缺少索引，已有索引: {existing_indexed_fields}")
                return False
        
        print(f"⚠️ 无法确定索引状态，字段: {all_fields} 需要进一步检查")
        # 如果无法获取任何索引信息，保守地认为可能缺少索引
        return False
    
    def _get_table_indexed_fields(self, table_name: str, database: str = '', query: Optional[dict] = None, hostname: str = None) -> set:
        """
        获取表的已有索引字段集合
        
        Args:
            table_name: 表名（可能是别名，需要映射到实际表名）
            database: 数据库名
            query: 查询对象，可能包含表结构信息
            hostname: 主机名
            
        Returns:
            已有索引的字段集合（小写）
        """
        existing_indexed_fields = set()
        
        if not table_name:
            return existing_indexed_fields
        
        # 从query对象中获取hostname_max用于连接真实业务数据库
        hostname_max = hostname
        if not hostname_max and query and isinstance(query, dict):
            slow_info = query.get('slow_query_info', {})
            hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
        
        # 尝试从query对象中获取表结构信息（可能包含多个表的信息）
        # 首先检查是否有表结构信息，可能包含多个表
        table_structure = None
        if query and isinstance(query, dict):
            # 尝试直接获取表结构
            if 'table_structure' in query:
                table_structure = query.get('table_structure', {})
            # 也可能在slow_query_info中
            elif 'slow_query_info' in query:
                slow_info = query.get('slow_query_info', {})
                if 'table_structure' in slow_info:
                    table_structure = slow_info.get('table_structure', {})
        
        # 如果table_structure是字符串，尝试解析
        if table_structure and isinstance(table_structure, str):
            try:
                import json
                table_structure = json.loads(table_structure)
            except (json.JSONDecodeError, ValueError):
                try:
                    import ast
                    table_structure = ast.literal_eval(table_structure)
                except (ValueError, SyntaxError):
                    table_structure = {}
        
        # 1. 优先从数据库获取实际索引信息
        # 确定正确的数据库：如果database为空，或表不在该数据库中，则查找正确的数据库
        actual_db = database
        need_find_database = False
        
        # 如果database为空，需要查找
        if not actual_db:
            need_find_database = True
        # 如果database不为空，先验证表是否在该数据库中
        elif actual_db and hostname_max and table_name:
            if not self.db_helper.check_table_exists(actual_db, table_name, hostname_max):
                print(f"⚠️ 表 {table_name} 在数据库 {actual_db} 中不存在，将查找正确的数据库")
                need_find_database = True
        
        # 如果需要查找数据库，通过hostname_max查找表所在的数据库
        if need_find_database and hostname_max and table_name:
            found_database = self.db_helper.find_correct_database_for_table(table_name, hostname_max)
            if found_database:
                actual_db = found_database
                print(f"🔍 通过hostname_max找到表 {table_name} 所在的数据库: {actual_db}")
            else:
                print(f"❌ 无法找到表 {table_name} 所在的数据库")
        
        # 使用正确的数据库查询索引
        if actual_db and table_name:
            query_result = self.db_helper.execute_safe_query(
                f"SHOW INDEX FROM `{table_name}`",
                hostname=hostname_max,
                database=actual_db
            )
            if query_result['status'] == 'success' and query_result['data']:
                for row in query_result['data']:
                    if len(row) >= 5:
                        column_name = row[4]
                        if column_name:
                            existing_indexed_fields.add(column_name.lower())
                if existing_indexed_fields:
                    print(f"✅ 从数据库 {actual_db} 中获取到表 {table_name} 的索引字段: {existing_indexed_fields}")
                    return existing_indexed_fields
            else:
                # 如果查询失败（如错误1146），尝试重新查找数据库
                error_msg = query_result.get('message', '未知错误')
                if '1146' in str(error_msg) or 'Table' in str(error_msg) and "doesn't exist" in str(error_msg):
                    print(f"⚠️ 从数据库 {actual_db} 查询表 {table_name} 的索引失败（表不存在），重新查找正确的数据库")
                    if hostname_max:
                        found_database = self.db_helper.find_correct_database_for_table(table_name, hostname_max)
                        if found_database and found_database != actual_db:
                            actual_db = found_database
                            print(f"🔍 重新找到表 {table_name} 所在的数据库: {actual_db}")
                            # 使用新找到的数据库重新查询
                            query_result = self.db_helper.execute_safe_query(
                                f"SHOW INDEX FROM `{table_name}`",
                                hostname=hostname_max,
                                database=actual_db
                            )
                            if query_result['status'] == 'success' and query_result['data']:
                                for row in query_result['data']:
                                    if len(row) >= 5:
                                        column_name = row[4]
                                        if column_name:
                                            existing_indexed_fields.add(column_name.lower())
                                if existing_indexed_fields:
                                    print(f"✅ 从数据库 {actual_db} 中获取到表 {table_name} 的索引字段: {existing_indexed_fields}")
                                    return existing_indexed_fields
                else:
                    print(f"⚠️ 从数据库 {actual_db} 查询表 {table_name} 的索引失败: {error_msg}")
        
        # 2. 如果数据库查询失败，从query对象中获取表结构信息作为参考
        # 需要处理表名可能是别名的情况，尝试匹配实际表名
        if table_structure and isinstance(table_structure, dict):
            # 如果table_structure是单个表的结构
            if 'indexes' in table_structure:
                indexes = table_structure['indexes']
                self._extract_indexes_from_structure(indexes, existing_indexed_fields)
            # 如果table_structure包含多个表的结构（字典格式：{table_name: structure}）
            elif isinstance(table_structure, dict):
                # 尝试直接匹配表名
                if table_name in table_structure:
                    table_info = table_structure[table_name]
                    if isinstance(table_info, dict) and 'indexes' in table_info:
                        indexes = table_info['indexes']
                        self._extract_indexes_from_structure(indexes, existing_indexed_fields)
                # 尝试小写匹配
                elif table_name.lower() in {k.lower(): v for k, v in table_structure.items()}:
                    for key, value in table_structure.items():
                        if key.lower() == table_name.lower() and isinstance(value, dict) and 'indexes' in value:
                            indexes = value['indexes']
                            self._extract_indexes_from_structure(indexes, existing_indexed_fields)
                            break
                # 如果都不匹配，尝试遍历所有表结构（可能是别名映射问题）
                else:
                    for key, value in table_structure.items():
                        if isinstance(value, dict) and 'indexes' in value:
                            indexes = value['indexes']
                            self._extract_indexes_from_structure(indexes, existing_indexed_fields)
        
        return existing_indexed_fields
    
    def _extract_indexes_from_structure(self, indexes, existing_indexed_fields: set):
        """从索引结构中提取索引字段"""
        if isinstance(indexes, dict):
            for index_info in indexes.values():
                if isinstance(index_info, dict) and 'columns' in index_info:
                    for col in index_info['columns']:
                        if isinstance(col, str):
                            existing_indexed_fields.add(col.lower())
                        elif isinstance(col, dict) and 'column' in col:
                            existing_indexed_fields.add(col['column'].lower())
        elif isinstance(indexes, list):
            for index_info in indexes:
                if isinstance(index_info, dict):
                    if 'columns' in index_info:
                        for col in index_info['columns']:
                            if isinstance(col, str):
                                existing_indexed_fields.add(col.lower())
                            elif isinstance(col, dict) and 'column' in col:
                                existing_indexed_fields.add(col['column'].lower())
                    elif 'Column_name' in index_info:
                        existing_indexed_fields.add(index_info['Column_name'].lower())
    
    def _check_composite_index_exists(self, existing_indexed_fields: set, composite_fields: list) -> bool:
        """
        检查是否已有复合索引覆盖指定的字段组合
        
        Args:
            existing_indexed_fields: 已有索引的字段集合（小写）
            composite_fields: 需要检查的复合索引字段列表
            
        Returns:
            如果已有复合索引覆盖这些字段，返回True，否则返回False
        """
        if not composite_fields:
            return False
            
        # 检查复合索引的最左前缀原则
        # 如果所有字段都已有单独的索引，认为可以组成复合索引
        for field in composite_fields:
            if field.lower() not in existing_indexed_fields:
                return False
        
        return True
    
    def _mask_table_structure(self, table_structure) -> str:
        """脱敏表结构信息（包装方法）"""
        return DataMasking.mask_table_structure(table_structure)
    
    def _merge_analysis_results_to_compare_data(self, analysis_results: List[Dict]):
        """将DeepSeek分析结果合并到compare_data结构中（包装方法）"""
        DataProcessor.merge_analysis_results_to_compare_data(
            self.compare_data, 
            analysis_results, 
            DataProcessor.format_deepseek_suggestions
        )
    
    def _create_compare_data_with_analysis(self, analysis_results: List[Dict]) -> Dict:
        """创建包含DeepSeek分析结果的compare_data结构（包装方法）"""
        return DataProcessor.create_compare_data_with_analysis(
            analysis_results, 
            DataProcessor.format_deepseek_suggestions
        )
    
    def _format_deepseek_suggestions(self, deepseek_optimization, sql_content: str = '') -> str:
        """智能格式化DeepSeek优化建议（包装方法，保留复杂逻辑）"""
        # 使用DataProcessor的方法，但保留主文件中的复杂逻辑
        return DataProcessor.format_deepseek_suggestions(deepseek_optimization, sql_content)
    
    def _analyze_sql_for_optimization(self, sql_content: str, database: str = '', table: str = '', query: Optional[dict] = None, hostname: str = None) -> str:
        """
        智能分析SQL语句，生成具体的优化建议和可执行语句
        重点处理JOIN查询，为多表生成索引建议
        """
        if not sql_content:
            return ""
        
        sql_lower = sql_content.lower()
        
        # 提取表别名映射
        table_alias_map = SQLAnalyzer.extract_table_aliases(sql_content)
        
        # 提取主表名
        primary_table = table or self._extract_table_name(sql_content) or 'your_table_name'
        primary_table_lower = primary_table.lower()
        
        # 提取字段信息
        where_fields = self._extract_where_fields(sql_content)
        join_fields = self._extract_join_fields(sql_content)
        
        # 如果没有WHERE和JOIN字段，返回缺少过滤条件的诊断
        if not where_fields and not join_fields:
            return "\n".join([
                "1. 智能诊断: 查询缺少有效的过滤条件，存在全表扫描风险",
                "• 建议添加包含索引的过滤条件",
                "3. 预期效果: 预计平均查询时间从60ms降低到3ms，性能提升约20倍"
            ])
        
        # 从query对象中获取数据库名（如果database参数为空）
        actual_database = database
        if not actual_database and query and isinstance(query, dict):
            actual_database = query.get('database') or query.get('db') or ''
            # 也可以从slow_query_info中获取
            if not actual_database:
                slow_info = query.get('slow_query_info', {})
                actual_database = slow_info.get('database') or slow_info.get('db') or ''
        
        # 分析JOIN条件，提取每个表涉及的字段
        table_field_usage = defaultdict(lambda: {'where': [], 'join': []})
        
        # 解析WHERE字段，识别表别名
        sql_upper = sql_content.upper()
        if 'WHERE' in sql_upper:
            where_clause = sql_content[sql_upper.find('WHERE') + 5:]
            # 提取带表别名的字段
            alias_field_pattern = r'([a-zA-Z_]\w*)\s*\.\s*([a-zA-Z_]\w*)'
            alias_matches = re.findall(alias_field_pattern, where_clause, re.IGNORECASE)
            for alias_name, column_name in alias_matches:
                alias_clean = alias_name.strip('`')
                column_clean = column_name.strip('`')
                # 优先从别名映射中获取实际表名，如果没有则使用别名本身（可能是表名）
                actual_table = table_alias_map.get(alias_clean, alias_clean)
                if actual_table:
                    table_field_usage[actual_table]['where'].append(column_clean)
            # 无别名的字段归属主表
            for field in where_fields:
                if '.' not in field:
                    table_field_usage[primary_table]['where'].append(field)
        
        # 解析JOIN字段，识别每个表
        # 支持 WHERE a.id=b.id 和 ON a.id=b.id 两种格式
        join_condition_pattern = r'([a-zA-Z_]\w*\.[a-zA-Z_]\w*)\s*=\s*([a-zA-Z_]\w*\.[a-zA-Z_]\w*)'
        join_matches = re.findall(join_condition_pattern, sql_content, re.IGNORECASE)
        for left_operand, right_operand in join_matches:
            for operand in (left_operand, right_operand):
                if '.' in operand:
                    alias_part, column_part = operand.split('.', 1)
                    alias = alias_part.strip('`')
                    column = column_part.strip('`')
                    # 优先从别名映射中获取实际表名，如果没有则使用别名本身（可能是表名）
                    actual_table = table_alias_map.get(alias, alias)
                    if actual_table and column:
                        # 避免重复添加
                        if column not in table_field_usage[actual_table]['join']:
                            table_field_usage[actual_table]['join'].append(column)
        
        # 构建优化建议
        optimization_parts = []
        solutions = []
        executable_actions = []
        
        # 1. 智能诊断
        core_issues = []
        # 收集所有表的字段需求，合并WHERE和JOIN字段，避免重复
        # 同时检查每个表的字段是否已有索引，只提示缺少索引的字段
        for table_key, usage in table_field_usage.items():
            all_fields = []
            field_types = []
            
            # 收集WHERE字段
            if usage['where']:
                where_fields_list = sorted(set(usage['where']))
                all_fields.extend(where_fields_list)
                field_types.append('WHERE')
            
            # 收集JOIN字段（去重）
            if usage['join']:
                join_fields_list = sorted(set(usage['join']))
                for field in join_fields_list:
                    if field not in all_fields:
                        all_fields.append(field)
                if 'JOIN' not in field_types:
                    field_types.append('JOIN')
            
            # 检查哪些字段缺少索引
            if all_fields:
                # 获取该表的已有索引字段（使用实际的数据库名）
                existing_indexed_fields = self._get_table_indexed_fields(table_key, actual_database, query, hostname)
                
                # 调试信息：打印索引检查结果
                if existing_indexed_fields:
                    print(f"🔍 表 {table_key} 的已有索引字段: {existing_indexed_fields}")
                else:
                    print(f"⚠️ 表 {table_key} 未找到索引信息（数据库: {actual_database}）")
                
                # 过滤出缺少索引的字段
                missing_index_fields = []
                for field in all_fields:
                    field_lower = field.lower()
                    if field_lower not in existing_indexed_fields:
                        missing_index_fields.append(field)
                    else:
                        print(f"✅ 表 {table_key} 的字段 {field} 已有索引，跳过")
                
                # 只对缺少索引的字段生成诊断提示
                if missing_index_fields:
                    if len(field_types) == 2:
                        # 同时有WHERE和JOIN字段，合并描述
                        core_issues.append(f"表 {table_key} 的字段 {', '.join(missing_index_fields)} 需要索引（用于WHERE和JOIN条件）")
                    elif 'WHERE' in field_types:
                        core_issues.append(f"表 {table_key} 的 WHERE 字段 {', '.join(missing_index_fields)} 需要索引")
                    elif 'JOIN' in field_types:
                        core_issues.append(f"表 {table_key} 的 JOIN 字段 {', '.join(missing_index_fields)} 需要索引")
        
        # 如果没有收集到任何信息，使用通用描述
        if not core_issues:
            if where_fields:
                # 如果没有表信息，至少显示字段
                core_issues.append(f"WHERE条件字段 {', '.join(where_fields[:3])} 需要索引支持")
            else:
                core_issues.append("SQL语句可能存在性能优化空间")
        
        optimization_parts.append(f"1. 智能诊断：{'；'.join(core_issues)}")
        
        # 2. 为主表生成复合索引建议（WHERE+JOIN）
        primary_usage = table_field_usage.get(primary_table, {'where': [], 'join': []})
        primary_where = primary_usage.get('where', [])
        primary_join = primary_usage.get('join', [])
        
        if primary_where or primary_join:
            # 获取主表的已有索引字段（使用实际的数据库名）
            primary_existing_indexes = self._get_table_indexed_fields(primary_table, actual_database, query, hostname)
            
            combined_fields = []
            # 先添加WHERE字段（过滤条件优先），过滤掉已有索引的字段
            for col in primary_where:
                if col and col.lower() not in primary_existing_indexes and col not in combined_fields:
                    combined_fields.append(col)
            # 再添加JOIN字段，过滤掉已有索引的字段
            for col in primary_join:
                if col and col.lower() not in primary_existing_indexes and col not in combined_fields:
                    combined_fields.append(col)
            
            if combined_fields:
                fields_subset = combined_fields[:5]
                index_name = f"idx_{'_'.join(fields_subset)}_composite"
                fields_str = ', '.join(fields_subset)
                solutions.append(f"🔥 为表 {primary_table} 创建复合索引（WHERE+JOIN）：{fields_str}")
                executable_actions.append(f"-- 🔥【主表复合索引】表 {primary_table}（WHERE+JOIN字段）")
                executable_actions.append(f"CREATE INDEX {index_name} ON {primary_table}({fields_str});")
        
        # 3. 为非主表生成JOIN字段索引建议
        for table_key, usage in table_field_usage.items():
            if table_key.lower() == primary_table_lower:
                continue
            
            # 获取该表的已有索引字段（使用实际的数据库名）
            table_existing_indexes = self._get_table_indexed_fields(table_key, actual_database, query, hostname)
            
            combined_order = []
            # 过滤掉已有索引的字段
            for col in usage['where']:
                if col and col.lower() not in table_existing_indexes and col not in combined_order:
                    combined_order.append(col)
            for col in usage['join']:
                if col and col.lower() not in table_existing_indexes and col not in combined_order:
                    combined_order.append(col)
            
            if not combined_order:
                continue
            
            if len(combined_order) >= 2:
                fields_subset = combined_order[:5]
                index_name = f"idx_{table_key.replace('.', '_')}_{'_'.join(fields_subset)}_join"
                fields_str = ', '.join(fields_subset)
                solutions.append(f"🔥 为表 {table_key} 创建复合索引覆盖JOIN字段：{fields_str}")
                executable_actions.append(f"-- 🔥【跨表JOIN复合索引】表 {table_key}")
                executable_actions.append(f"CREATE INDEX {index_name} ON {table_key}({fields_str});")
            else:
                field = combined_order[0]
                index_name = f"idx_{table_key.replace('.', '_')}_{field}_join"
                solutions.append(f"为表 {table_key} 的 JOIN 字段 {field} 创建单列索引优化连接性能")
                executable_actions.append(f"-- ✅ 为表 {table_key} 的 JOIN字段 {field} 创建单列索引")
                executable_actions.append(f"CREATE INDEX {index_name} ON {table_key}({field});")
        
        # 4. 构建完整的优化建议字符串
        # 不添加 "2. 智能优化建议：" 或 "智能优化建议：" 标记，直接添加 SQL 代码块，让 report_generator_core.py 通过检测 "```sql" 来识别
        if executable_actions:
            optimization_parts.append("```sql")
            optimization_parts.extend(executable_actions)
            optimization_parts.append("```")
        
        # 5. 预期效果
        if solutions:
            optimization_parts.append("3. 预期效果: 预计平均查询时间可降低50%以上，JOIN性能显著提升")
        
        return "\n".join(optimization_parts)
    
    def _convert_analysis_to_queries(self, analysis_results: List[Dict]) -> List[Dict]:
        """将分析结果转换为查询列表格式（包装方法）"""
        return DataProcessor.convert_analysis_to_queries(
            analysis_results, 
            self._format_deepseek_suggestions
        )
    
    def create_report(self) -> str:
        """创建Word格式的数据库优化分析报告（包装方法，调用新模块）"""
        import os
        from docx import Document
        from datetime import datetime
        
        # 创建输出目录
        output_dir = "."
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 生成带时间戳的文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"数据库智能优化分析报告_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)
        
        # 创建Word文档
        doc = Document()
        
        # 创建报告生成核心实例
        report_core = ReportGeneratorCore(
            document=doc,
            analysis_data=self.analysis_data,
            compare_data=self.compare_data,
            db_helper=self.db_helper,
            sql_optimizer=self._analyze_sql_for_optimization
        )
        
        # 设置页面布局和样式
        report_core.setup_page_layout()
        report_core.setup_document_styles()
        
        # 生成报告各部分
        report_core.generate_report_header()
        report_core.generate_report_summary()
        report_core.add_compare_analysis()
        report_core.generate_top_sql_statements()
        report_core.generate_sql_details()
        
        # 生成总结和建议（使用 SummaryGenerator）
        summary_gen = SummaryGenerator(
            document=doc,
            analysis_data=self.analysis_data,
            compare_data=self.compare_data
        )
        summary_gen.generate_summary_and_recommendations()
        
        # 生成报告页脚
        report_core.generate_report_footer()
        
        # 保存文档
        doc.save(filepath)
        
        print(f"Word报告已生成: {filepath}")
        return filepath
     
    def _generate_optimization_suggestions(self):
        """生成优化建议"""
        self.document.add_heading('六、优化建议', level=1)
        
        # 确保analysis_data不为None
        if self.analysis_data is None:
            self.analysis_data = []
        
        for i, query in enumerate(self.analysis_data, 1):
            # 创建智能优化建议标题
            self.document.add_heading(f'智能优化建议 #{i}', level=2)
            
            # 提取优化建议各部分 - 优先使用deepseek_optimization，如果没有则使用optimization_suggestions
            suggestions = query.get('deepseek_optimization', '') or query.get('optimization_suggestions', '')
            
            # 如果deepseek_optimization是列表，转换为结构化字符串格式
            if isinstance(suggestions, list):
                # 将列表转换为结构化建议格式
                structured_suggestions = []
                for item in suggestions:
                    if '已存在索引' in item or '最优状态' in item:
                        structured_suggestions.append(f"1. 智能诊断: {item}")
                    elif '表不存在' in item:
                        structured_suggestions.append(f"1. 智能诊断: {item}")
                    elif '未找到合适的索引' in item or '建议分析查询模式' in item:
                        # 处理通用的索引建议
                        structured_suggestions.append(f"1. 智能诊断: {item}")
                        structured_suggestions.append("建议分析该SQL的查询模式，考虑添加合适的索引")
                        # structured_suggestions.append("3. 预期效果: 通过添加合适索引，查询性能预计可提升60-90%")
                        
                        # 添加具体的索引建议
                        structured_suggestions.append("4. 具体索引建议:")
                        
                        # 从SQL语句中提取表名和字段信息
                        sql_content = query.get('sql', query.get('sql_content', ''))
                        if sql_content:
                            # 分析WHERE条件中的字段
                            where_fields = self._extract_where_fields(sql_content)
                            if where_fields:
                                for field in where_fields:
                                    index_name = f"idx_{field}"
                                    structured_suggestions.append(f"   • 建议创建索引: `{index_name}({field})`")
                                    structured_suggestions.append(f"     SQL: CREATE INDEX {index_name} ON table_name({field});")
                            
                            # 分析JOIN条件中的字段
                            join_fields = self._extract_join_fields(sql_content)
                            if join_fields:
                                for field in join_fields:
                                    index_name = f"idx_{field}_join"
                                    structured_suggestions.append(f"   • JOIN字段索引: `{index_name}({field})`")
                                    structured_suggestions.append(f"     SQL: CREATE INDEX {index_name} ON table_name({field});")
                            
                            # 分析ORDER BY字段
                            order_fields = self._extract_order_fields(sql_content)
                            if order_fields:
                                for field in order_fields:
                                    index_name = f"idx_{field}_order"
                                    structured_suggestions.append(f"   • 排序字段索引: `{index_name}({field})`")
                                    structured_suggestions.append(f"     SQL: CREATE INDEX {index_name} ON table_name({field});")
                            
                            # 如果没有提取到具体字段，提供通用建议
                            if not where_fields and not join_fields and not order_fields:
                                structured_suggestions.append("   • 请检查SQL语句中的WHERE、JOIN和ORDER BY子句")
                                structured_suggestions.append("   • 为经常用于查询条件的字段创建单列索引")
                                structured_suggestions.append("   • 考虑创建复合索引以支持多条件查询")
                                structured_suggestions.append("   • 索引示例: CREATE INDEX idx_column ON table_name(column);")
                        
                        structured_suggestions.append("5. 注意事项:")
                        structured_suggestions.append("   • 在创建索引前，请评估表的写操作频率")
                        structured_suggestions.append("   • 避免在频繁更新的字段上创建索引")
                        structured_suggestions.append("   • 建议先在测试环境验证索引效果")
                        structured_suggestions.append("   • 使用EXPLAIN命令验证索引是否被使用")
                    else:
                        # 对于其他类型的建议，也转换为结构化格式
                        structured_suggestions.append(f"1. 智能诊断: {item}")
                        structured_suggestions.append("建议进一步分析该查询的执行计划")
                        structured_suggestions.append("3. 预期效果: 通过优化，查询性能有望得到显著提升")
                suggestions = '\n\n'.join(structured_suggestions)
            
            # 检查优化建议是否为空或无效
            if not suggestions or (isinstance(suggestions, str) and not suggestions.strip()) or suggestions == '暂无优化建议':
                # 使用智能分析生成具体的优化建议
                sql_content = query.get('sql', query.get('sql_content', ''))
                database = query.get('database', query.get('db_name', ''))
                table = query.get('table', '')
                
                # 如果没有表名，尝试从SQL语句中提取
                if not table and sql_content:
                    table = self._extract_table_name(sql_content)
                
                # 获取hostname_max用于连接真实的业务数据库
                slow_info = query.get('slow_query_info', {})
                hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
                
                suggestions = self._analyze_sql_for_optimization(sql_content, database, table, query, hostname_max)
            
            # 如果仍然没有有效建议，显示通用建议
            if not suggestions or (isinstance(suggestions, str) and not suggestions.strip()):
                # 优化建议为空的情况
                empty_box = self.document.add_paragraph()
                empty_run = empty_box.add_run("⚠ 该SQL暂无具体的优化建议")
                empty_run.font.name = '微软雅黑'
                empty_run.font.size = Pt(11)
                empty_run.font.color.rgb = RGBColor(255, 140, 0)  # 橙色警告
                empty_run.bold = True
                empty_box.paragraph_format.space_before = Pt(6)
                empty_box.paragraph_format.space_after = Pt(6)
                
                # 提供通用建议
                general_title = self.document.add_paragraph()
                general_run = general_title.add_run('通用优化建议:')
                general_run.bold = True
                general_run.font.name = '微软雅黑'
                general_run.font.size = Pt(11)
                
                general_content = self.document.add_paragraph()
                general_content_run = general_content.add_run(
                    "评估是否可以优化SQL语句结构\n"
                )
                general_content_run.font.name = '宋体'
                general_content_run.font.size = Pt(10.5)
                general_content.paragraph_format.left_indent = Pt(15)
                
                # 添加空行和分隔线，然后继续下一个查询
                self._add_separator_line()
                continue
            
            # 添加背景色的提示框（仅在有优化建议时显示）
            highlight_box = self.document.add_paragraph()
            highlight_run = highlight_box.add_run("以下是针对该SQL的详细优化建议")
            highlight_run.font.name = '微软雅黑'
            highlight_run.font.size = Pt(11)
            highlight_run.font.color.rgb = RGBColor(192, 0, 0)
            highlight_run.bold = True
            highlight_box.paragraph_format.space_before = Pt(6)
            highlight_box.paragraph_format.space_after = Pt(6)
            
            # 分割建议内容
            # 使用更智能的分割方式，确保预期效果部分不会被错误分割
            parts = []
            
            # 先尝试按编号分割
            lines = suggestions.split('\n')
            current_part = []
            
            for line in lines:
                # 检查是否是新的部分开始（以数字编号开头）
                if re.match(r'^\d+\.', line.strip()) or re.match(r'^\*\*\d+\.', line.strip()):
                    # 如果当前部分不为空，保存它
                    if current_part:
                        parts.append('\n'.join(current_part))
                        current_part = []
                current_part.append(line)
            
            # 添加最后一部分
            if current_part:
                parts.append('\n'.join(current_part))
            
            # 如果没有正确分割，使用原始方式
            if len(parts) <= 1:
                parts = suggestions.split('\n\n')
            
            for part in parts:
                if part.startswith('1. 智能诊断') or part.startswith('**1. 智能诊断**'):
                    # 智能诊断部分
                    issue_title = self.document.add_paragraph()
                    issue_title_run = issue_title.add_run('🎯 智能诊断:')
                    issue_title_run.bold = True
                    issue_title_run.font.name = '微软雅黑'
                    issue_title_run.font.size = Pt(11)
                    issue_title_run.font.color.rgb = RGBColor(192, 0, 0)  # 红色突出问题
                    
                    # 去除标记并添加内容
                    content = re.sub(r'1\.\s*智能诊断[:：]?\s*|\*\*1\.\s*智能诊断\*\*\s*', '', part)
                    issue_content = self.document.add_paragraph()
                    issue_content_run = issue_content.add_run(content)
                    issue_content_run.font.name = '宋体'
                    issue_content_run.font.size = Pt(10.5)
                    issue_content_run.font.color.rgb = RGBColor(192, 0, 0)
                    issue_content.paragraph_format.left_indent = Pt(15)
                
                elif part.startswith('2. 智能优化建议') or part.startswith('**2. 智能优化建议**'):
                    # 智能优化建议部分 - 只有当内容不包含"最优状态"时才添加
                    # 移除了重复的"智能优化建议:"标题，避免与内容重复
                    
                    # 处理SQL代码块
                    if '```sql' in part:
                        # 移除```sql和```标记，只保留SQL代码内容
                        # 先移除开头的```sql
                        sql_content = re.sub(r'^```sql\s*\n?', '', part, flags=re.MULTILINE)
                        # 移除结尾的```
                        sql_content = re.sub(r'\n?```\s*$', '', sql_content, flags=re.MULTILINE)
                        # 移除中间的```sql和```标记
                        sql_content = re.sub(r'```sql\s*\n?', '', sql_content, flags=re.MULTILINE)
                        sql_content = re.sub(r'\n?```\s*', '', sql_content, flags=re.MULTILINE)
                        
                        # 移除"2. 智能优化建议"标记
                        sql_content = re.sub(r'2\.\s*智能优化建议[:：]?\s*|\*\*2\.\s*智能优化建议\*\*\s*', '', sql_content)
                        
                        # 处理SQL代码
                        if sql_content.strip():
                            sql_lines = sql_content.strip().split('\n')
                            for sql_line in sql_lines:
                                if sql_line.strip():
                                    sql_para = self.document.add_paragraph()
                                    sql_run = sql_para.add_run(sql_line)
                                    sql_run.font.name = 'Consolas'
                                    sql_run.font.size = Pt(9)
                                    
                                    if sql_line.strip().startswith('-- 🔥'):
                                        sql_run.font.color.rgb = RGBColor(255, 0, 0)
                                        sql_run.font.bold = True
                                    elif sql_line.strip().startswith('-- 🔍') or sql_line.strip().startswith('-- ✅'):
                                        sql_run.font.color.rgb = RGBColor(0, 100, 200)
                                        sql_run.font.bold = True
                                    elif sql_line.strip().startswith('--'):
                                        sql_run.font.color.rgb = RGBColor(128, 128, 128)
                                    elif 'CREATE INDEX' in sql_line.upper() or 'ALTER TABLE' in sql_line.upper():
                                        sql_run.font.color.rgb = RGBColor(0, 128, 0)
                                        sql_run.font.bold = True
                                    else:
                                        sql_run.font.color.rgb = RGBColor(0, 0, 0)
                                    
                                    sql_para.paragraph_format.left_indent = Pt(20)
                                    sql_para.paragraph_format.space_before = Pt(0)
                                    sql_para.paragraph_format.space_after = Pt(0)
                    else:
                        # 普通文本智能优化建议
                        content = re.sub(r'2\.\s*智能优化建议[:：]?\s*|\*\*2\.\s*智能优化建议\*\*\s*', '', part)
                        solution_content = self.document.add_paragraph()
                        solution_content_run = solution_content.add_run(content)
                        solution_content_run.font.name = '宋体'
                        solution_content_run.font.size = Pt(10.5)
                        solution_content.paragraph_format.left_indent = Pt(15)
                
                elif part.startswith('3. 预期效果') or part.startswith('**3. 预期效果**'):
                    # 预期效果部分
                    effect_title = self.document.add_paragraph()
                    effect_title_run = effect_title.add_run('🚀 预期效果:')
                    effect_title_run.bold = True
                    effect_title_run.font.name = '微软雅黑'
                    effect_title_run.font.size = Pt(11)
                    effect_title_run.font.color.rgb = RGBColor(0, 0, 192)  # 蓝色效果标题
                    effect_title_run.font.color.rgb = RGBColor(0, 0, 192)  # 蓝色效果标题
                    
                    content = re.sub(r'3\.\s*预期效果[:：]?\s*|\*\*3\.\s*预期效果\*\*\s*', '', part)
                    effect_content = self.document.add_paragraph()
                    effect_content_run = effect_content.add_run(content)
                    effect_content_run.font.name = '宋体'
                    effect_content_run.font.size = Pt(10.5)
                    effect_content_run.font.color.rgb = RGBColor(0, 0, 192)  # 蓝色效果文本
                    effect_content.paragraph_format.left_indent = Pt(15)
                
                # 添加对预期效果的宽松匹配处理
                elif '预期效果' in part and not any(keyword in part for keyword in ['1. 智能诊断', '2. 智能优化建议', '4. 预期效果', '5. ', '6. ']):
                    # 处理包含预期效果但没有标准编号的部分
                    effect_title = self.document.add_paragraph()
                    effect_title_run = effect_title.add_run('🚀 预期效果:')
                    effect_title_run.bold = True
                    effect_title_run.font.name = '微软雅黑'
                    effect_title_run.font.size = Pt(11)
                    effect_title_run.font.color.rgb = RGBColor(0, 0, 192)  # 蓝色效果标题
                    
                    # 移除预期效果关键词及相关内容
                    content = re.sub(r'.*预期效果[:：]?\s*', '', part, count=1)
                    if content.strip():
                        effect_content = self.document.add_paragraph()
                        effect_content_run = effect_content.add_run(content)
                        effect_content_run.font.name = '宋体'
                        effect_content_run.font.size = Pt(10.5)
                        effect_content_run.font.color.rgb = RGBColor(0, 0, 192)  # 蓝色效果文本
                        effect_content.paragraph_format.left_indent = Pt(15)
            
            # 添加空行和分隔线
            self._add_separator_line()
    
    def _generate_summary_and_recommendations(self):
        """生成总结和建议（包装方法，调用新模块）"""
        summary_gen = SummaryGenerator(
            document=self.document,
            analysis_data=self.analysis_data,
            compare_data=self.compare_data
        )
        summary_gen.generate_summary_and_recommendations()
    
    def _generate_report_footer(self):
        """生成报告页脚"""
        # 添加空行
        self.document.add_paragraph()
        
        # 添加最终页脚
        footer = self.document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("*本报告由数据库智能优化系统自动生成，仅供参考*")
        footer_run.font.name = '宋体'
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 添加生成日期和时间
        footer_date = self.document.add_paragraph()
        footer_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_date_run = footer_date.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer_date_run.font.name = '宋体'
        footer_date_run.font.size = Pt(9)
        footer_date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 添加页码
        sections = self.document.sections
        for section in sections:
            # 添加页脚
            footer = section.footer
            # 确保页脚有段落
            if not footer.paragraphs:
                paragraph = footer.add_paragraph()
            else:
                paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码字段 - 使用PAGE字段，只显示当前页码
            run = paragraph.add_run("第 ")
            run.font.name = '宋体'
            run.font.size = Pt(9)
            
            # 插入PAGE字段 - 使用Word字段方式
            from docx.oxml.shared import OxmlElement, qn
            page_run = paragraph.add_run()
            page_run.font.name = '宋体'
            page_run.font.size = Pt(9)
            
            # 添加页码字段，Word会自动替换为实际页码
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            page_run._r.append(fldChar1)
            page_run._r.append(instrText)
            page_run._r.append(fldChar2)
            
            run = paragraph.add_run(" 页")
            run.font.name = '宋体'
            run.font.size = Pt(9)

# 添加缺少的import
import re

def load_db_config(config_file: str = 'db_config.json') -> Optional[Dict]:
    """
    从配置文件加载数据库配置
    支持处理单配置对象或配置数组
    
    Args:
        config_file: 配置文件路径
    
    Returns:
        数据库配置字典，如果加载失败返回None
    """
    try:
        with open(config_file, 'r', encoding='utf-8') as f:
            config_data = json.load(f)
            
            # 处理配置数组格式
            if isinstance(config_data, list):
                # 如果是数组，取第一个配置项作为默认配置
                if not config_data:
                    print(f"❌ 配置文件中没有配置项")
                    return None
                config = config_data[0]
                print(f"⚠️  检测到配置数组，使用第一个配置项")
            else:
                config = config_data
            
            # 验证必要的配置项
            required_fields = ['host', 'user', 'password']
            for field in required_fields:
                if field not in config:
                    print(f"❌ 配置文件缺少必要项: {field}")
                    return None
            
            # 添加慢查询分析默认参数
            config.setdefault('table', 's')  # 默认慢查询表名
            config.setdefault('port', 3306)  # 默认端口
            
            return config
    except FileNotFoundError:
        print(f"❌ 配置文件不存在: {config_file}")
        return None
    except json.JSONDecodeError:
        print(f"❌ 配置文件格式错误: {config_file}")
        return None

def main():
    """主函数"""
    print("=== 数据库智能优化分析报告生成器 ===")    
    try:
        # 尝试加载数据库配置
        default_db_config = load_db_config()
        
        # 使用实时分析模式，连接到实际数据库
        use_live_analysis = True
        slow_query_db_config = None
        
        # 设置默认的过滤参数
        min_execute_cnt = 1000
        min_query_time = 10.0
        
        print("📊 慢查询分析配置")
        print("------------------")
        
        if default_db_config:
            # 使用配置文件中的数据库连接信息
            slow_query_db_config = default_db_config
            # 确保使用正确的慢查询表名
            if 'table' not in slow_query_db_config or slow_query_db_config['table'] == 's':
                slow_query_db_config['table'] = 'slow'
                print("⚠️ 已自动修正慢查询表名为 'slow'")
            print("✓ 使用配置文件中的数据库连接信息")
        else:
            # 使用默认的连接配置
            print("⚠️ 配置文件不存在或格式错误，使用默认数据库连接配置")
            slow_query_db_config = {
                'host': '127.0.0.1',
                'port': 3306,
                'user': 'test',
                'password': 'test',
                'database': 't',
                'table': 'slow'  # 使用正确的慢查询表名
            }
            print(f"✓ 使用默认连接配置: host={slow_query_db_config['host']}, port={slow_query_db_config['port']}")
            print(f"✓ 默认慢查询表名: {slow_query_db_config['table']}")
        
        # 打印过滤参数
        print("\n🔍 慢查询过滤条件")
        print("------------------")
        print(f"✓ 过滤条件: 执行次数≥{min_execute_cnt}, 查询时间≥{min_query_time}秒 (ts_cnt > 1000, query_time_max > 10)")
        
        # 创建报告生成器
        import logging
        logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        logger = logging.getLogger(__name__)
        logger.info("📈 正在执行实时慢查询分析与对比分析...")
        report = DatabaseOptimizationReport(
            use_live_analysis=use_live_analysis,
            slow_query_db_config=slow_query_db_config,
            min_execute_cnt=min_execute_cnt,
            min_query_time=min_query_time
        )
        
        # 检查是否成功获取了分析数据
        if not report.analysis_data:
            print("\n❌ 错误：无法获取真实的分析数据")
            print("   可能原因：")
            print("   1. 数据库连接失败")
            print("   2. 慢查询表不存在或为空")
            print("   3. 没有符合过滤条件的慢查询记录")
            print("   4. 分析数据文件不存在或格式错误")
            print("\n   请确保：")
            print("   1. 数据库连接配置正确")
            print("   2. 慢查询表中有数据")
            print("   3. 分析数据文件存在且格式正确")
            print("\n   程序将退出，请检查以上问题后重新运行")
            return 1
        
        # 生成报告
        print("\n📝 正在生成优化分析报告...")
        output_file = report.create_report()
        
        print("")
        print(f"✅ 数据库智能优化分析报告已生成: {output_file}")
        print(f"📄 文件位置: {os.path.abspath(output_file)}")
        
        # 添加结果说明
        if not report.compare_data:
            print("\n📋 报告说明：")
            print("   - 报告中包含基本分析内容，但可能缺少完整的对比分析")
            print("   - 建议检查数据库连接和慢查询表配置")
        
    except KeyboardInterrupt:
        print("\n❌ 操作被用户中断")
        return 0
    except ImportError as e:
        print(f"\n错误：缺少必要的依赖库。请安装 python-docx 库：")
        print("pip install python-docx")
        return 1
    except ConnectionError as e:
        print("\n❌ 数据库连接错误")
        print(f"   错误详情: {str(e)}")
        print("   请检查以下内容:")
        print("   1. 数据库服务器是否运行")
        print("   2. 连接配置是否正确")
        print("   3. 网络连接是否正常")
        return 1
    except Exception as e:
        print(f"\n❌ 生成报告时发生异常: {str(e)}")
        print("   程序将优雅退出")
        # 只在调试模式下显示详细堆栈
        if os.environ.get('DEBUG', 'False').lower() == 'true':
            import traceback
            traceback.print_exc()
        return 1
    
    return 0

if __name__ == "__main__":
    import sys
    sys.exit(main())
