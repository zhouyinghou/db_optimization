#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告生成核心模块
包含所有与 create_report 相关的报告生成方法
"""

import re
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_BREAK
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls

from data_masking import DataMasking
from sql_analyzer import SQLAnalyzer
from database_helper import DatabaseHelper
from summary_generator import SummaryGenerator

# 尝试导入智能优化建议模块（可选）
try:
    from intelligent_optimization_suggestions import IntelligentOptimizationSuggestions
    INTELLIGENT_OPTIMIZER_AVAILABLE = True
except ImportError:
    INTELLIGENT_OPTIMIZER_AVAILABLE = False
    IntelligentOptimizationSuggestions = None


class ReportGeneratorCore:
    """报告生成核心类，包含所有报告生成相关的方法"""
    
    def __init__(self, document: Document, analysis_data: List[Dict], 
                 compare_data: Optional[Dict] = None, db_helper: DatabaseHelper = None,
                 sql_optimizer=None):
        """
        初始化报告生成核心类
        
        Args:
            document: Word文档对象
            analysis_data: 分析数据
            compare_data: 对比数据
            db_helper: 数据库辅助类实例
        """
        self.document = document
        self.analysis_data = analysis_data
        self.compare_data = compare_data
        self.db_helper = db_helper
        self.sql_optimizer = sql_optimizer
        self.intelligent_optimizer = None
        
        # 仅当未传入sql_optimizer时，才启用新的智能优化建议生成器作为兜底
        if self.sql_optimizer is None and INTELLIGENT_OPTIMIZER_AVAILABLE and IntelligentOptimizationSuggestions:
            try:
                self.intelligent_optimizer = IntelligentOptimizationSuggestions(
                    db_helper=self.db_helper
                )
            except Exception:
                self.intelligent_optimizer = None
    
    def setup_page_layout(self):
        """设置页面布局"""
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
    
    def setup_document_styles(self):
        """设置文档样式"""
        styles = self.document.styles
        
        # 标题1样式 - 黑体
        title_style = styles['Heading 1']
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_font.size = Pt(16)
        title_font.bold = True
        title_font.color.rgb = RGBColor(31, 73, 125)
        title_para_format = title_style.paragraph_format
        title_para_format.space_before = Pt(6)
        title_para_format.space_after = Pt(6)
        
        # 标题2样式 - 楷体
        title2_style = styles['Heading 2']
        title2_font = title2_style.font
        title2_font.name = 'Times New Roman'
        title2_font._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        title2_font.size = Pt(14)
        title2_font.bold = True
        title2_font.color.rgb = RGBColor(31, 73, 125)
        title2_para_format = title2_style.paragraph_format
        title2_para_format.space_before = Pt(4)
        title2_para_format.space_after = Pt(4)
        
        # 标题3样式
        title3_style = styles['Heading 3']
        title3_font = title3_style.font
        title3_font.name = 'Times New Roman'
        title3_font._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        title3_font.size = Pt(12)
        title3_font.bold = True
        title3_para_format = title3_style.paragraph_format
        title3_para_format.space_before = Pt(2)
        title3_para_format.space_after = Pt(2)
        
        # 正文样式 - 三号字体
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        normal_font.size = Pt(16)
        normal_para_format = normal_style.paragraph_format
        normal_para_format.space_after = Pt(3)
        normal_para_format.line_spacing = 1.0
        normal_para_format.left_indent = Pt(0)
    
    def add_separator_line(self):
        """添加分隔线"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.LINE)
        
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '366092')
        pBdr.append(bottom)
    
    def generate_report_header(self):
        """生成报告标题和头部信息"""
        title = self.document.add_heading('数据库智能优化分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 73, 125)
        title_run.font.underline = False
        
        date_info = self.document.add_paragraph()
        date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        current_year = datetime.now().strftime('%Y')
        current_month = datetime.now().strftime('%m')
        current_day = datetime.now().strftime('%d')
        current_time = datetime.now().strftime('%H:%M:%S')
        date_run = date_info.add_run(f"生成日期: {current_year}年{current_month}月{current_day}日 {current_time}")
        date_run.font.name = '宋体'
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(64, 64, 64)
        
        mask_notice = self.document.add_paragraph()
        mask_notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mask_run = mask_notice.add_run("⚠️ 本报告已对敏感信息（库名、IP、表名等）进行脱敏处理")
        mask_run.font.name = '微软雅黑'
        mask_run.font.size = Pt(12)
        mask_run.font.color.rgb = RGBColor(192, 0, 0)
        mask_run.bold = True
        
        self.add_separator_line()
    
    def add_compare_analysis(self):
        """添加上个月与上上个月的慢查询对比分析"""
        # 直接在摘要下面添加分隔线
        self.add_separator_line()
        
        # 添加标题 - 左对齐并添加序列号
        title = self.document.add_heading('二、慢查询对比分析', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        try:
            # 添加错误处理和数据验证
            if not self.compare_data:
                para = self.document.add_paragraph()
                para.add_run("无法获取对比分析数据，可能原因：").font.size = Pt(12)
                para.add_run("\n1. 数据库连接失败")
                para.add_run("\n2. 没有足够的慢查询数据")
                return
            
            # 安全获取对比数据
            compare_data = self.compare_data
            
            # 获取月份信息（如果可用）
            last_month_name = compare_data.get('last_month', {}).get('name', '当前月')
            previous_month_name = compare_data.get('previous_month', {}).get('name', '上月')
            
            # 添加子标题
            sub_title = self.document.add_heading('（一）慢查询同比', level=2)
            sub_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            para = self.document.add_paragraph()
            para.add_run(f"对比期间: {previous_month_name} vs {last_month_name}").font.size = Pt(12)
            para.paragraph_format.space_after = Pt(18)
            
            # 添加总体对比表格
            comparison_table = self.document.add_table(rows=1, cols=4)
            comparison_table.style = 'Table Grid'
            
            # 表头
            hdr_cells = comparison_table.rows[0].cells
            hdr_cells[0].text = '指标'
            hdr_cells[1].text = previous_month_name
            hdr_cells[2].text = last_month_name
            hdr_cells[3].text = '变化率'
            
            # 设置表头样式
            for cell in hdr_cells:
                cell_run = cell.paragraphs[0].runs[0]
                cell_run.bold = True
                cell_run.font.name = '微软雅黑'
                cell_run.font.size = Pt(11)
                cell_run.font.color.rgb = RGBColor(255, 255, 255)
                
                # 设置表头背景色
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "366092")
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # 居中对齐
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 安全获取数据，使用默认值避免KeyError
            prev_total_count = str(compare_data.get('previous_month', {}).get('total_count', 0))
            last_total_count = str(compare_data.get('last_month', {}).get('total_count', 0))
            count_change = compare_data.get('comparison', {}).get('count_change', 0)
            
            # 添加数据行（仅保留慢查询总数）
            rows_data = [
                ['慢查询总数', prev_total_count, last_total_count, f"{count_change:.2f}%↑"]
            ]
            
            for row_data in rows_data:
                row_cells = comparison_table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = cell_data
                    # 设置单元格样式
                    cell_run = row_cells[i].paragraphs[0].runs[0]
                    cell_run.font.name = '宋体'
                    cell_run.font.size = Pt(10.5)  # 与摘要表格字体大小一致
                    
                    # 数据列居中
                    if i > 0:
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 添加分析说明
            self.document.add_paragraph('')
            analysis_para = self.document.add_paragraph()
            analysis_para.add_run('分析说明：').bold = True
            
            # 生成分析内容（仅保留慢查询数量分析）
            analysis_text = []
            if count_change > 0:
                analysis_text.append(f"1. {last_month_name}慢查询数量较{previous_month_name}增加了{count_change:.2f}%，系统性能有所下降")
            elif count_change < 0:
                analysis_text.append(f"1. {last_month_name}慢查询数量较{previous_month_name}减少了{abs(count_change):.2f}%，系统性能有所改善")
            else:
                analysis_text.append(f"1. 两个月的慢查询数量保持不变")
            
            # 添加分析文本
            for text in analysis_text:
                analysis_content_para = self.document.add_paragraph()
                analysis_content_run = analysis_content_para.add_run(text)
                analysis_content_run.font.name = '宋体'
                analysis_content_run.font.size = Pt(16)  # 三号字体
                # 设置段落缩进，与摘要正文保持一致
                analysis_content_para.paragraph_format.left_indent = Pt(0)
            
            # 添加建议
            self.document.add_paragraph()
            suggestion_para = self.document.add_paragraph()
            suggestion_para.add_run('改进建议：').bold = True
            suggestion_para.add_run('\n')
            
            # 生成建议
            if count_change > 0:
                suggestion_text = suggestion_para.add_run('1. 建议重点关注新增的慢查询，分析其产生原因。\n')
                suggestion_text.font.size = Pt(16)  # 三号字体
                suggestion_text = suggestion_para.add_run('2. 检查是否有新增的查询模式或数据量增长导致慢查询增加。\n')
                suggestion_text.font.size = Pt(16)  # 三号字体
                suggestion_text = suggestion_para.add_run('3. 考虑对频繁访问的表进行索引优化或查询重写。')
                suggestion_text.font.size = Pt(16)  # 三号字体
            elif count_change < 0:
                suggestion_text = suggestion_para.add_run('1. 慢查询数量有所减少，继续保持当前的优化策略。\n')
                suggestion_text.font.size = Pt(16)  # 三号字体
                suggestion_text = suggestion_para.add_run('2. 定期检查系统性能，确保优化效果持续。\n')
                suggestion_text.font.size = Pt(16)  # 三号字体
                suggestion_text = suggestion_para.add_run('3. 考虑预防性优化措施，避免性能退化。')
                suggestion_text.font.size = Pt(16)  # 三号字体
            else:
                suggestion_text = suggestion_para.add_run('1. 慢查询数量保持稳定，继续监控系统性能。\n')
                suggestion_text.font.size = Pt(16)  # 三号字体
                suggestion_text = suggestion_para.add_run('2. 定期检查新增查询的性能影响。\n')
                suggestion_text.font.size = Pt(16)  # 三号字体
                suggestion_text = suggestion_para.add_run('3. 考虑预防性优化措施，避免性能退化。')
                suggestion_text.font.size = Pt(16)  # 三号字体
          
        except Exception as e:
            # 捕获所有异常，确保报告生成不会中断
            error_para = self.document.add_paragraph()
            error_para.add_run(f"生成对比分析时发生错误: {str(e)}").font.color.rgb = RGBColor(255, 0, 0)
            error_para.add_run("\n将继续生成报告的其他部分...")
    
    def generate_report_summary(self):
        """生成报告摘要"""
        self.document.add_heading('一、摘要', level=1)
        
        # 报告概述 - 第一行空两格
        summary = self.document.add_paragraph()
        summary_run = summary.add_run("  本报告基于MySQL慢查询日志分析，提供了数据库性能问题诊断和优化建议。报告包含了对慢查询SQL的详细分析，识别了性能瓶颈，并提供了针对性的智能优化建议。")
        summary_run.font.name = '宋体'
        summary_run.font.size = Pt(16)  # 三号字体
        
        # 获取上个月的数据
        last_month_queries = []
        if self.compare_data and 'last_month' in self.compare_data and 'queries' in self.compare_data['last_month']:
            last_month_queries = self.compare_data['last_month']['queries']
        else:
            # 如果没有明确的上个月数据，使用所有分析数据
            last_month_queries = self.analysis_data
        
        # 确保last_month_queries不为None
        if last_month_queries is None:
            last_month_queries = []
        
        total_queries = len(last_month_queries)

        # 分析统计信息（仅上个月）
        total_queries = len(last_month_queries)
        
        # 使用更美观的表格样式
        stats_table = self.document.add_table(rows=1, cols=3)
        stats_table.style = 'Table Grid'
        
        # 设置表格宽度
        for cell in stats_table.rows[0].cells:
            cell.width = Inches(2.5)
        
        # 表头
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = '统计项'
        hdr_cells[1].text = '数值'
        hdr_cells[2].text = '说明'
        
        # 设置表头样式
        for cell in hdr_cells:
            cell_run = cell.paragraphs[0].runs[0]
            cell_run.bold = True
            cell_run.font.name = '微软雅黑'
            cell_run.font.size = Pt(11)
            cell_run.font.color.rgb = RGBColor(255, 255, 255)
            cell_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置表头背景色
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "366092")
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # 居中对齐
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 获取上个月的实际年份和月份名称
        last_month_name = '上月完整月份'
        if self.compare_data and 'last_month' in self.compare_data:
            last_month_name = self.compare_data['last_month'].get('name', '上月完整月份')
        else:
            # 如果没有具体名称，计算上个月的实际年份和月份
            from datetime import timedelta
            today = datetime.now()
            last_month = today.replace(day=1) - timedelta(days=1)
            last_month_name = f'{last_month.year}年{last_month.month}月'
        
        # 添加数据行（添加上月时间范围和阈值说明）
        data_rows = [
            ('慢查询总数', str(total_queries), '上月符合条件的慢查询数量'),
            ('筛选时间范围', last_month_name, '基于慢查询日志时间戳筛选'),
            ('执行次数阈值', '≥1000次', '仅分析执行次数达到1000次及以上的慢查询'),
            ('查询时间阈值', '≥10秒', '仅分析查询时间达到10秒及以上的慢查询')
        ]
        
        for i, (item, value, desc) in enumerate(data_rows):
            row_cells = stats_table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = value
            row_cells[2].text = desc
            
            # 设置数据行样式
            for j, cell in enumerate(row_cells):
                cell_run = cell.paragraphs[0].runs[0]
                cell_run.font.name = '宋体'
                cell_run.font.size = Pt(10.5)
                
                # 居中对齐
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 交替行背景色
                if i % 2 == 1:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "F2F2F2")
                    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def get_sorted_queries(self):
        """获取按执行次数降序、平均时间降序、数据库名排序的查询列表"""
        # 获取上个月的数据
        last_month_queries = []
        if hasattr(self, 'compare_data') and self.compare_data and 'last_month' in self.compare_data:
            last_month_queries = self.compare_data['last_month'].get('queries', [])
        
        # 如果没有上个月的特定数据，则使用所有分析数据
        if not last_month_queries:
            last_month_queries = self.analysis_data
        
        # 确保last_month_queries不为None
        if last_month_queries is None:
            last_month_queries = []
        
        # 对查询进行排序 - 按执行次数降序、平均时间降序、数据库名排序
        try:
            sorted_queries = sorted(last_month_queries, 
                                   key=lambda x: (
                                       int(x.get('slow_query_info', {}).get('execute_cnt', 0)),
                                       float(x.get('slow_query_info', {}).get('query_time_max') or 
                                             x.get('slow_query_info', {}).get('query_time') or 
                                             x.get('query_time', 0)),
                                       x.get('slow_query_info', {}).get('db_name') or x.get('db_name', '')
                                   ), 
                                   reverse=True)
        except (TypeError, ValueError):
            # 如果排序失败，使用原始顺序
            sorted_queries = last_month_queries
        
        return sorted_queries
    
    def generate_top_sql_statements(self):
        """生成Top SQL语句列表（仅显示上个月数据）"""
        self.document.add_heading('三、性能问题SQL概览', level=1)
        
        # 添加简介
        intro = self.document.add_paragraph()
        intro_run = intro.add_run("下表展示了按照执行次数降序、平均时间降序、数据库名排序的上个月慢查询SQL概览，帮助快速识别影响系统性能的关键SQL语句。")
        intro_run.font.name = '宋体'
        intro_run.font.size = Pt(10.5)
        
        # 获取排序后的查询列表
        sorted_queries = self.get_sorted_queries()

        # 创建表格，使用更美观的样式
        sql_table = self.document.add_table(rows=1, cols=6)
        sql_table.style = 'Table Grid'
        
        # 设置表格列宽自适应
        sql_table.autofit = True
        
        # 设置各列的初始宽度（根据内容类型设置合理宽度）
        sql_table.columns[0].width = Inches(0.5)   # 排名：窄列
        sql_table.columns[1].width = Inches(2.5)   # SQLID：较宽（显示SQL片段）
        sql_table.columns[2].width = Inches(1.2)   # 数据库：中等
        sql_table.columns[3].width = Inches(1.0)   # 表名：中等
        sql_table.columns[4].width = Inches(0.8)   # 执行次数：较窄
        sql_table.columns[5].width = Inches(1.0)   # 平均时间：较窄
        
        # 表头
        hdr_cells = sql_table.rows[0].cells
        hdr_cells[0].text = '排名'
        hdr_cells[1].text = 'SQLID'
        hdr_cells[2].text = '数据库'
        hdr_cells[3].text = '表名'
        hdr_cells[4].text = '执行次数'
        hdr_cells[5].text = '平均时间(ms)'
        
        # 设置表头样式
        for cell in hdr_cells:
            cell_run = cell.paragraphs[0].runs[0]
            cell_run.bold = True
            cell_run.font.name = '微软雅黑'
            cell_run.font.size = Pt(11)
            cell_run.font.color.rgb = RGBColor(255, 255, 255)
            
            # 设置表头背景色
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "366092")
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # 居中对齐
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加数据行
        for i, query in enumerate(sorted_queries[:10], 1):  # 只显示前10个
            # 获取SQL内容，兼容不同的字段名
            sql_content = query.get('sql', query.get('sql_content', ''))
            # 对SQL内容进行脱敏处理
            masked_sql_content = DataMasking.mask_sql(sql_content)
            sql_id = masked_sql_content[:32] + '...' if len(masked_sql_content) > 32 else masked_sql_content
            
            # 尝试从SQL语句中提取表名
            table_name = SQLAnalyzer.extract_table_name(sql_content)
            
            row_cells = sql_table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = sql_id
            # 兼容两种数据结构：slow_query_info对象或直接字段
            # 优先使用slow_query_info对象，如果没有则直接使用顶层字段
            slow_info = query.get('slow_query_info', {})
            db_name = slow_info.get('db_name') or query.get('db_name', '未知')
            
            # 对数据库名进行脱敏处理
            db_name = DataMasking.mask_db_name(db_name)
            
            # 如果数据库名是默认值或未知，尝试通过表名查找正确的数据库
            # 使用hostname_max连接真实的业务数据库
            if db_name in ['未知', 'db', 't'] and table_name and self.db_helper:
                # 获取hostname_max作为真实的业务数据库IP
                hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
                correct_db = self.db_helper.find_correct_database_for_table(table_name, hostname_max)
                if correct_db:
                    db_name = correct_db
                    # 对找到的数据库名进行脱敏处理
                    db_name = DataMasking.mask_db_name(db_name)
                else:
                    # 如果找不到数据库，标记为库表未找到
                    db_name = '库表未找到'
            
            # 如果表名为空，标记为库表未找到
            if not table_name:
                table_name = '库表未找到'
            else:
                # 对表名进行脱敏处理
                table_name = DataMasking.mask_table_name(table_name)
            
            execute_cnt = slow_info.get('execute_cnt') or query.get('execute_cnt', 0)
            # 优先使用query_time_max，其次是query_time
            query_time = slow_info.get('query_time_max') or slow_info.get('query_time') or query.get('query_time', 0)
            
            row_cells[2].text = str(db_name)
            row_cells[3].text = str(table_name)
            row_cells[4].text = str(execute_cnt)
            # 显示查询时间（毫秒）
            row_cells[5].text = f"{query_time}ms"
            
            # 设置数据行样式
            for cell in row_cells:
                cell_run = cell.paragraphs[0].runs[0]
                cell_run.font.name = '宋体'
                cell_run.font.size = Pt(10.5)
                
                # 居中对齐
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 交替行背景色
                if i % 2 == 1:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "F2F2F2")
                    cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # 添加分隔线
        self.add_separator_line()
    
    def generate_report_footer(self):
        """生成报告页脚"""
        # 添加空行
        self.document.add_paragraph()
        
        # 添加最终页脚
        footer = self.document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("*本报告由数据库智能优化系统自动生成，仅供参考*")
        footer_run.font.name = '宋体'
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 添加生成日期和时间
        footer_date = self.document.add_paragraph()
        footer_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_date_run = footer_date.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer_date_run.font.name = '宋体'
        footer_date_run.font.size = Pt(9)
        footer_date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 添加页码
        sections = self.document.sections
        for section in sections:
            # 添加页脚
            footer = section.footer
            # 确保页脚有段落
            if not footer.paragraphs:
                paragraph = footer.add_paragraph()
            else:
                paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码字段 - 使用PAGE字段，只显示当前页码
            run = paragraph.add_run("第 ")
            run.font.name = '宋体'
            run.font.size = Pt(9)
            
            # 插入PAGE字段 - 使用Word字段方式
            page_run = paragraph.add_run()
            page_run.font.name = '宋体'
            page_run.font.size = Pt(9)
            
            # 添加页码字段，Word会自动替换为实际页码
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            
            page_run._r.append(fldChar1)
            page_run._r.append(instrText)
            page_run._r.append(fldChar2)
            
            run = paragraph.add_run(" 页")
            run.font.name = '宋体'
            run.font.size = Pt(9)
    
    def check_composite_index_exists(self, existing_indexed_fields: set, composite_fields: list) -> bool:
        """
        检查是否已有复合索引覆盖指定的字段组合
        
        Args:
            existing_indexed_fields: 已有索引的字段集合（小写）
            composite_fields: 需要检查的复合索引字段列表
            
        Returns:
            如果已有复合索引覆盖这些字段，返回True，否则返回False
        """
        if not composite_fields:
            return False
            
        # 检查复合索引的最左前缀原则
        # 如果所有字段都已有单独的索引，认为可以组成复合索引
        for field in composite_fields:
            if field.lower() not in existing_indexed_fields:
                return False
        
        return True
    
    def check_indexes_exist(self, database: str, table_name: str, where_fields: list, join_fields: list, order_by_fields: list, query: Optional[dict] = None) -> bool:
        """
        检查所有相关字段是否都有索引
        
        Args:
            database: 数据库名
            table_name: 表名
            where_fields: WHERE条件字段列表
            join_fields: JOIN条件字段列表
            order_by_fields: ORDER BY字段列表
            query: 查询对象，考虑JSON中的表结构信息作为参考
            
        Returns:
            如果所有字段都有索引，返回True，否则返回False
        """
        if not table_name:
            return False
        
        # 🎯 关键修复：如果提供了query参数且包含表结构信息，则跳过表存在性检查
        # 避免在没有数据库连接的情况下返回False
        if query and isinstance(query, dict) and 'table_structure' in query:
            print(f"ℹ️ 使用query参数中的表结构信息，跳过表存在性检查")
        elif database and table_name and self.db_helper:
            # 从query对象中获取hostname_max用于连接真实业务数据库
            hostname_max = None
            if query and isinstance(query, dict):
                slow_info = query.get('slow_query_info', {})
                hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
            
            if not self.db_helper.check_table_exists(database, table_name, hostname_max):
                print(f"⚠️ 表 {table_name} 在数据库 {database} 中不存在，无法检查索引")
                return False
        
        # 检查所有相关字段
        all_fields = set()
        all_fields.update([f.lower() for f in where_fields])
        all_fields.update([f.lower() for f in join_fields])
        all_fields.update([f.lower() for f in order_by_fields])
        
        if not all_fields:
            return False
        
        # 🔥 关键修复：优先从数据库读取真实索引信息，如果数据库查询失败，则从JSON数据中参考
        existing_indexed_fields = set()
        database_query_successful = False
        
        # 1. 优先从数据库获取实际索引信息（使用hostname_max连接真实业务数据库）
        # 从query对象或hostname参数中获取hostname_max
        hostname_max = None
        if query and isinstance(query, dict):
            slow_info = query.get('slow_query_info', {})
            hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
        
        if database and table_name and self.db_helper:
            # 使用execute_safe_query直接查询索引信息（支持hostname参数）
            query_result = self.db_helper.execute_safe_query(
                f"SHOW INDEX FROM `{table_name}`",
                hostname=hostname_max,
                database=database
            )
            if query_result['status'] == 'success' and query_result['data']:
                # 数据库查询成功且有数据
                for row in query_result['data']:
                    if len(row) >= 5:
                        column_name = row[4]
                        if column_name:
                            existing_indexed_fields.add(column_name.lower())
                if existing_indexed_fields:
                    database_query_successful = True
                    print(f"📊 从数据库读取到的索引字段: {existing_indexed_fields}")
        
        # 2. 如果数据库查询失败，从query对象中获取表结构信息作为参考
        if not database_query_successful and query and isinstance(query, dict) and 'table_structure' in query:
            table_structure = query.get('table_structure', {})
            # 如果table_structure是字符串，尝试解析
            if isinstance(table_structure, str):
                try:
                    import json
                    table_structure = json.loads(table_structure)
                except (json.JSONDecodeError, ValueError):
                    # 如果JSON解析失败，尝试使用ast.literal_eval（Python字符串表示）
                    try:
                        import ast
                        table_structure = ast.literal_eval(table_structure)
                    except (ValueError, SyntaxError):
                        table_structure = {}
            
            if table_structure and isinstance(table_structure, dict) and 'indexes' in table_structure:
                indexes = table_structure['indexes']
                
                # indexes可能是字典{index_name: index_info}或列表[index_info]
                if isinstance(indexes, dict):
                    # 字典格式：遍历values
                    for index_info in indexes.values():
                        if isinstance(index_info, dict) and 'columns' in index_info:
                            for col in index_info['columns']:
                                existing_indexed_fields.add(col.lower())
                elif isinstance(indexes, list):
                    # 列表格式
                    for index_info in indexes:
                        if isinstance(index_info, dict):
                            # 支持多种索引格式
                            if 'columns' in index_info:
                                # 格式1: {'columns': ['id']}
                                for col in index_info['columns']:
                                    existing_indexed_fields.add(col.lower())
                            elif 'Column_name' in index_info:
                                # 格式2: {'Column_name': 'id'} (MySQL SHOW INDEXES格式)
                                existing_indexed_fields.add(index_info['Column_name'].lower())
                
                if existing_indexed_fields:
                    print(f"📋 从JSON数据中参考到的索引字段: {existing_indexed_fields}")
        
        # 3. 检查所有字段是否都有索引
        if existing_indexed_fields:
            # 检查是否所有字段都有单独的索引
            all_fields_have_individual_indexes = True
            fields_without_individual_indexes = []
            
            for field in all_fields:
                if field not in existing_indexed_fields:
                    all_fields_have_individual_indexes = False
                    fields_without_individual_indexes.append(field)
            
            if all_fields_have_individual_indexes:
                # 所有字段都有单独索引，检查是否需要复合索引
                # 如果WHERE条件中有多个字段，建议复合索引
                if len(where_fields) > 1:
                    print(f"ℹ️ 所有字段都有单独索引，但WHERE条件中有多个字段，建议复合索引")
                    return False  # 返回False表示建议创建复合索引
                else:
                    print(f"✅ 所有字段都有索引，字段: {all_fields}, 已有索引: {existing_indexed_fields}")
                    return True
            else:
                print(f"❌ 字段 {fields_without_individual_indexes} 缺少索引，已有索引: {existing_indexed_fields}")
                return False
        
        print(f"⚠️ 无法确定索引状态，字段: {all_fields} 需要进一步检查")
        # 如果无法获取任何索引信息，保守地认为可能缺少索引
        return False
    
    def generate_sql_details(self):
        """生成SQL详细信息"""
        self.document.add_heading('四、SQL详细分析', level=1)
        
        # 获取排序后的查询列表，与"三、性能问题SQL概览"保持一致
        sorted_queries = self.get_sorted_queries()
        
        # 只显示前10个SQL的详细分析
        for i, query in enumerate(sorted_queries[:10], 1):
            self.document.add_heading(f'SQL #{i}', level=2)
            
            # SQL基本信息
            sql_info_title = self.document.add_paragraph()
            sql_info_title_run = sql_info_title.add_run('🔍 SQL语句:')
            sql_info_title_run.bold = True
            sql_info_title_run.font.name = '微软雅黑'
            sql_info_title_run.font.size = Pt(11)
            sql_info_title_run.font.color.rgb = RGBColor(31, 73, 125)  # 深蓝色标题
            
            # SQL代码块美化，安全访问sql字段
            sql_para = self.document.add_paragraph()
            sql_content = query.get('sql', query.get('sql_content', '未知SQL'))
            
            # 尝试提取表名，如果query中没有table字段
            # 需要在SQL脱敏之前提取表名，避免从脱敏后的SQL中提取到错误的表名
            table_name = query.get('table')
            if not table_name:
                table_name = SQLAnalyzer.extract_table_name(sql_content)
            
            # 对SQL内容进行脱敏处理
            masked_sql_content = DataMasking.mask_sql(sql_content)
            sql_run = sql_para.add_run(masked_sql_content)
            sql_run.font.name = 'Consolas'
            sql_run.font.size = Pt(9)
            
            # 设置代码块样式
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "F5F5F5")
            sql_para._p.get_or_add_pPr().append(shading_elm)
            sql_para.paragraph_format.left_indent = Pt(15)
            sql_para.paragraph_format.space_before = Pt(6)
            sql_para.paragraph_format.space_after = Pt(6)
            
            # 执行信息
            info_table = self.document.add_table(rows=1, cols=2)
            info_table.style = 'Table Grid'
            
            # 设置表格宽度
            info_table.columns[0].width = Inches(1.5)
            info_table.columns[1].width = Inches(4.0)
            
            hdr_cells = info_table.rows[0].cells
            hdr_cells[0].text = '属性'
            hdr_cells[1].text = '值'
            
            # 设置表头样式
            for cell in hdr_cells:
                cell_run = cell.paragraphs[0].runs[0]
                cell_run.bold = True
                cell_run.font.name = '微软雅黑'
                cell_run.font.size = Pt(11)
                cell_run.font.color.rgb = RGBColor(255, 255, 255)
                
                # 设置表头背景色
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "366092")
                cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # 居中对齐
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 兼容两种数据结构：slow_query_info对象或直接字段
            slow_info = query.get('slow_query_info', {})
            host_ip = slow_info.get('ip') or query.get('hostname_max') or query.get('ip', '未知')
            # 对IP地址进行脱敏处理
            host_ip = DataMasking.mask_ip(host_ip)
            
            # 优先使用slow_query_info中的数据，如果没有则使用顶层字段
            db_name = slow_info.get('db_name') or query.get('db_name', '未知')
            
            # 对数据库名进行脱敏处理
            db_name = DataMasking.mask_db_name(db_name)
            execute_cnt = slow_info.get('execute_cnt') or query.get('execute_cnt', '0')
            query_time = slow_info.get('query_time') or query.get('query_time', 0.0)
            
            # 如果数据库名是默认值或未知，尝试通过表名查找正确的数据库
            # 使用hostname_max连接真实的业务数据库
            if db_name in ['未知', 'db', 't'] and table_name and self.db_helper:
                # 获取hostname_max作为真实的业务数据库IP
                hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
                correct_db = self.db_helper.find_correct_database_for_table(table_name, hostname_max)
                if correct_db:
                    db_name = correct_db
                    # 对找到的数据库名进行脱敏处理
                    db_name = DataMasking.mask_db_name(db_name)
                else:
                    # 如果找不到数据库，标记为库表未找到
                    db_name = '库表未找到'
            
            # 如果表名为空，标记为库表未找到
            if not table_name:
                table_name = '库表未找到'
            else:
                # 对表名进行脱敏处理
                table_name = DataMasking.mask_table_name(table_name)
            
            info_rows = [
                ('数据库', db_name),
                ('主机IP', host_ip),
                ('表名', table_name),
                ('执行次数', str(execute_cnt)),
                ('平均查询时间', f"{query_time}ms")
            ]
            
            for i_row, (prop, value) in enumerate(info_rows):
                row_cells = info_table.add_row().cells
                row_cells[0].text = prop
                # 确保值是字符串类型
                row_cells[1].text = str(value)
                
                # 设置属性列样式
                prop_cell_run = row_cells[0].paragraphs[0].runs[0]
                prop_cell_run.font.name = '微软雅黑'
                prop_cell_run.font.size = Pt(10.5)
                prop_cell_run.font.bold = True
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 设置值列样式
                value_cell_run = row_cells[1].paragraphs[0].runs[0]
                value_cell_run.font.name = '宋体'
                value_cell_run.font.size = Pt(10.5)
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 交替行背景色
                if i_row % 2 == 1:
                    for cell in row_cells:
                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "F2F2F2")
                        cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # 🔍 直接在此处添加优化建议 - 紧跟SQL语句之后
            # 使用原始SQL进行分析，确保字段名不受脱敏影响
            original_sql = query.get('sql', query.get('sql_content', sql_content))
            original_table_name = query.get('table') or SQLAnalyzer.extract_table_name(original_sql) or table_name
            self.add_optimization_suggestion_for_query(query, original_sql, original_table_name or 'unknown', i)
            
            # 添加分隔线
            self.add_separator_line()
    
    def add_optimization_suggestion_for_query(self, query: dict, sql_content: str, table_name: str, index: int):
        """为单个查询添加优化建议 - 直接跟在SQL语句后面"""
        
        # 首先尝试从当前查询中获取DeepSeek优化建议
        suggestions = query.get('deepseek_optimization', '') or query.get('optimization_suggestions', '')
        
        # 如果当前查询中没有DeepSeek建议，尝试从compare_data中查找对应的分析结果
        
        # 通过SQL内容匹配查找对应的分析结果
        analysis_queries = []  # 初始化分析查询列表
        
        if not suggestions and hasattr(self, 'compare_data') and self.compare_data:
            # 查找匹配的SQL分析结果
            if 'last_month' in self.compare_data and 'queries' in self.compare_data['last_month']:
                analysis_queries.extend(self.compare_data['last_month']['queries'])
            if 'previous_month' in self.compare_data and 'queries' in self.compare_data['previous_month']:
                analysis_queries.extend(self.compare_data['previous_month']['queries'])
        
        # 使用集合来避免重复处理相同的SQL语句
        processed_sqls = set()
        
        for i, analysis_query in enumerate(analysis_queries):
            analysis_sql = analysis_query.get('sql', '').strip()
            
            # 跳过空SQL或已处理的SQL
            if not analysis_sql or analysis_sql in processed_sqls:
                continue
                
            processed_sqls.add(analysis_sql)
            
            # 使用模糊匹配而不是精确匹配
            if analysis_sql == sql_content.strip() or \
               (analysis_sql in sql_content.strip() or sql_content.strip() in analysis_sql):
                suggestions = analysis_query.get('deepseek_optimization', '') or analysis_query.get('optimization_suggestions', '')
                if suggestions:
                    break
        
        # 获取hostname_max用于连接真实的业务数据库
        hostname_max = None
        if isinstance(query, dict):
            slow_info = query.get('slow_query_info', {})
            hostname_max = slow_info.get('hostname_max') or slow_info.get('ip') or query.get('hostname_max') or query.get('ip')
        
        # 如果deepseek_optimization是列表，转换为结构化字符串格式
        if isinstance(suggestions, list):
            # 直接调用智能分析函数生成具体的可执行SQL语句
            database = query.get('database', query.get('db_name', '')) if isinstance(query, dict) else ''
            # 确保传递原始表名信息
            original_table = query.get('table') if isinstance(query, dict) else None
            suggestions = self.analyze_sql_for_optimization(sql_content, database, original_table or table_name, query, hostname_max)
        else:
            # 对于字符串格式的建议，如果内容不够具体，也调用智能分析
            if not suggestions or suggestions == '暂无优化建议' or '建议分析查询模式' in suggestions:
                database = query.get('database', query.get('db_name', '')) if isinstance(query, dict) else ''
                # 确保传递原始表名信息
                original_table = query.get('table') if isinstance(query, dict) else None
                suggestions = self.analyze_sql_for_optimization(sql_content, database, original_table or table_name, query, hostname_max)
        
        # 检查优化建议是否为空或无效
        if not suggestions or (isinstance(suggestions, str) and not suggestions.strip()) or suggestions == '暂无优化建议':
            # 使用智能分析生成具体的优化建议
            database = query.get('database', query.get('db_name', '')) if isinstance(query, dict) else ''
            # 确保传递原始表名信息
            original_table = query.get('table') if isinstance(query, dict) else None
            suggestions = self.analyze_sql_for_optimization(sql_content, database, original_table or table_name, query, hostname_max)
        
        # 如果仍然没有有效建议，显示通用建议
        if not suggestions or (isinstance(suggestions, str) and not suggestions.strip()):
            # 添加通用优化建议标题
            general_title = self.document.add_paragraph()
            general_run = general_title.add_run('智能优化建议:')
            general_run.bold = True
            general_run.font.name = '微软雅黑'
            general_run.font.size = Pt(11)
            general_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色标题
            
            general_content = self.document.add_paragraph()
            general_content_run = general_content.add_run(
                "评估是否可以优化SQL语句结构\n"
            )
            general_content_run.font.name = '宋体'
            general_content_run.font.size = Pt(10.5)
            general_content.paragraph_format.left_indent = Pt(15)
            return
        
        self._render_structured_suggestions(suggestions)

    def _render_structured_suggestions(self, suggestions: str):
        """渲染包含“智能诊断/智能优化建议/预期效果”的结构化文本"""
        if not suggestions:
            return
        
        parts = []
        
        # 匹配1. 智能诊断（支持多种格式）
        diagnosis_match = re.search(r'(1\.\s*智能诊断[:：]?[^\n]*\n[^\n]*|智能诊断[:：][^\n]*)', suggestions)
        if diagnosis_match:
            diagnosis_content = diagnosis_match.group(0)
            if not diagnosis_content.startswith('1.'):
                diagnosis_content = "1. " + diagnosis_content
            parts.append(diagnosis_content)
        else:
            loose_diagnosis_match = re.search(r'(智能诊断[:：].*?)(?=智能优化建议|预期效果|$)', suggestions, re.DOTALL)
            if loose_diagnosis_match:
                diagnosis_content = loose_diagnosis_match.group(0).strip()
                if diagnosis_content.startswith('智能诊断：'):
                    diagnosis_content = diagnosis_content[5:]
                elif diagnosis_content.startswith('智能诊断:'):
                    diagnosis_content = diagnosis_content[4:]
                diagnosis_content = "1. 智能诊断:\n" + diagnosis_content.strip()
                parts.append(diagnosis_content)
        
        # 匹配2. 智能优化建议（支持多种格式，包含完整的```sql代码块）
        optimization_match = re.search(r'(2\.\s*智能优化建议.*?```sql.*?```)', suggestions, re.DOTALL)
        if optimization_match:
            parts.append(optimization_match.group(0))
        else:
            optimization_match = re.search(r'(2\.\s*智能优化建议[:：]?.*?)(?=\n\n[34]\.|预期效果|$)', suggestions, re.DOTALL)
            if optimization_match:
                parts.append(optimization_match.group(0))
            else:
                loose_optimization_match = re.search(r'(智能优化建议[:：].*?)(?=预期效果|$)', suggestions, re.DOTALL)
                if loose_optimization_match:
                    optimization_content = loose_optimization_match.group(0).strip()
                    if optimization_content.startswith('智能优化建议：'):
                        optimization_content = optimization_content[6:]
                    elif optimization_content.startswith('智能优化建议:'):
                        optimization_content = optimization_content[5:]
                    optimization_content = "2. 智能优化建议:\n" + optimization_content.strip()
                    parts.append(optimization_content)
        
        # 匹配3/4. 预期效果（支持多种格式）
        effect_match = re.search(r'([34]\.\s*[^\n]*预期效果[^\n]*[:：]?.*?)(?=\n\n[45]\.|$)', suggestions, re.DOTALL)
        if effect_match:
            parts.append(effect_match.group(0))
        else:
            general_effect_match = re.search(r'(.*?预期效果[:：].*?)($|\n\n)', suggestions, re.DOTALL)
            if general_effect_match:
                effect_content = general_effect_match.group(1).strip()
                if effect_content.startswith('预期效果：'):
                    effect_content = effect_content[5:]
                elif effect_content.startswith('预期效果:'):
                    effect_content = effect_content[4:]
                effect_content = "3. 预期效果:\n" + effect_content.strip()
                parts.append(effect_content)
        
        # 重新排序部分：确保智能诊断 -> 智能优化建议 -> 预期效果 的顺序
        reordered_parts = []
        diagnosis_part = None
        optimization_part = None
        effect_part = None
        
        for part in parts:
            if '智能诊断' in part and ('1.' in part or part.startswith('**1.') or part.startswith('智能诊断')):
                diagnosis_part = part
            elif '智能优化建议' in part and ('2.' in part or part.startswith('**2.') or part.startswith('智能优化建议')):
                optimization_part = part
            elif '预期效果' in part:
                effect_part = part
            else:
                reordered_parts.append(part)
        
        if diagnosis_part:
            reordered_parts.append(diagnosis_part)
        if optimization_part:
            reordered_parts.append(optimization_part)
        if effect_part:
            reordered_parts.append(effect_part)
        
        parts = reordered_parts
        
        for part in parts:
            if part.startswith('1. 智能诊断') or part.startswith('**1. 智能诊断**') or '智能诊断' in part:
                issue_title = self.document.add_paragraph()
                issue_title.paragraph_format.space_before = Pt(0)
                issue_title_run = issue_title.add_run('🎯 智能诊断:')
                issue_title_run.bold = True
                issue_title_run.font.name = '微软雅黑'
                issue_title_run.font.size = Pt(11)
                issue_title_run.font.color.rgb = RGBColor(192, 0, 0)
                
                content = re.sub(r'^1\.\s*智能诊断[:：]?\s*|^\*\*1\.\s*智能诊断\*\*\s*|^智能诊断[:：]?\s*', '', part)
                issue_content = self.document.add_paragraph()
                issue_content.paragraph_format.space_before = Pt(0)
                issue_content.paragraph_format.space_after = Pt(0)
                issue_content_run = issue_content.add_run(content)
                issue_content_run.font.name = '宋体'
                issue_content_run.font.size = Pt(10.5)
                issue_content_run.font.color.rgb = RGBColor(192, 0, 0)
                issue_content.paragraph_format.left_indent = Pt(15)
            
            elif part.startswith('2. 智能优化建议') or part.startswith('**2. 智能优化建议**') or '智能优化建议' in part:
                if not (part.strip().startswith('智能优化建议：') or part.strip().startswith('智能优化建议:')):
                    solution_title = self.document.add_paragraph()
                    solution_title.paragraph_format.space_before = Pt(0)
                    solution_title_run = solution_title.add_run('💡 智能优化建议:')
                    solution_title_run.bold = True
                    solution_title_run.font.name = '微软雅黑'
                    solution_title_run.font.size = Pt(11)
                    solution_title_run.font.color.rgb = RGBColor(0, 128, 0)
                
                if '```sql' in part:
                    sql_parts = part.split('```sql')
                    for sql_code_part in sql_parts[1:]:
                        if '```' in sql_code_part:
                            sql_code = sql_code_part.split('```')[0].strip()
                            if sql_code:
                                sql_lines = sql_code.split('\n')
                                for sql_line in sql_lines:
                                    if sql_line.strip():
                                        line_para = self.document.add_paragraph()
                                        line_run = line_para.add_run(sql_line)
                                        line_run.font.name = 'Consolas'
                                        line_run.font.size = Pt(9)
                                        
                                        if sql_line.strip().startswith('-- 🔥'):
                                            line_run.font.color.rgb = RGBColor(255, 0, 0)
                                            line_run.font.bold = True
                                        elif sql_line.strip().startswith('-- 🔍') or sql_line.strip().startswith('-- ✅'):
                                            line_run.font.color.rgb = RGBColor(0, 100, 200)
                                            line_run.font.bold = True
                                        elif sql_line.strip().startswith('-- 智能优化建议:'):
                                            line_run.font.color.rgb = RGBColor(0, 128, 0)
                                            line_run.font.bold = True
                                        elif sql_line.strip().startswith('--'):
                                            line_run.font.color.rgb = RGBColor(128, 128, 128)
                                        elif 'CREATE INDEX' in sql_line.upper() or 'ALTER TABLE' in sql_line.upper():
                                            line_run.font.color.rgb = RGBColor(0, 128, 0)
                                            line_run.font.bold = True
                                        elif 'EXPLAIN' in sql_line.upper() or 'SHOW' in sql_line.upper() or 'ANALYZE' in sql_line.upper():
                                            line_run.font.color.rgb = RGBColor(0, 100, 200)
                                        else:
                                            line_run.font.color.rgb = RGBColor(0, 0, 0)
                                        
                                        line_para.paragraph_format.left_indent = Pt(20)
                                        line_para.paragraph_format.space_before = Pt(0)
                                        line_para.paragraph_format.space_after = Pt(0)
                else:
                    content = re.sub(r'^2\.\s*智能优化建议[:：]?\s*|^\*\*2\.\s*智能优化建议\*\*\s*|^智能优化建议[:：]?\s*', '', part)
                    if content.strip():
                        solution_content = self.document.add_paragraph()
                        solution_content.paragraph_format.space_before = Pt(0)
                        solution_content_run = solution_content.add_run(content)
                        solution_content_run.font.name = '宋体'
                        solution_content_run.font.size = Pt(10.5)
                        solution_content.paragraph_format.left_indent = Pt(15)
            
            elif part.startswith('3. 预期效果') or part.startswith('**3. 预期效果**') or '预期效果' in part:
                effect_title = self.document.add_paragraph()
                effect_title_run = effect_title.add_run('🚀 预期效果:')
                effect_title_run.bold = True
                effect_title_run.font.name = '微软雅黑'
                effect_title_run.font.size = Pt(11)
                effect_title_run.font.color.rgb = RGBColor(0, 0, 192)
                
                content = re.sub(r'^3\.\s*预期效果[:：]?\s*|^\*\*3\.\s*预期效果\*\*\s*|^预期效果[:：]?\s*', '', part)
                effect_content = self.document.add_paragraph()
                effect_content.paragraph_format.space_before = Pt(0)
                effect_content.paragraph_format.space_after = Pt(0)
                effect_content_run = effect_content.add_run(content)
                effect_content_run.font.name = '宋体'
                effect_content_run.font.size = Pt(10.5)
                effect_content_run.font.color.rgb = RGBColor(0, 0, 192)
                effect_content.paragraph_format.left_indent = Pt(15)
    
    def analyze_sql_for_optimization(self, sql_content: str, database: str = '', table: str = '', query: Optional[dict] = None, hostname: str = None) -> str:
        """智能分析SQL语句，生成具体的优化建议和可执行语句"""
        if not sql_content:
            return ""
        
        # 优先使用调用方传入的优化器（保留拆分前的完整逻辑）
        if self.sql_optimizer:
            try:
                result = self.sql_optimizer(sql_content, database, table, query, hostname)
                if result and result.strip():
                    return result
            except Exception:
                pass
        
        # 如果没有传入表名，尝试从SQL中提取
        if not table:
            table = SQLAnalyzer.extract_table_name(sql_content) or ''

        # 尝试使用智能优化建议生成器
        if getattr(self, 'intelligent_optimizer', None):
            try:
                comprehensive_suggestions = self.intelligent_optimizer.generate_comprehensive_suggestions(
                    sql_content=sql_content,
                    database=database,
                    table=table,
                    query=query,
                    hostname=hostname
                )
                if comprehensive_suggestions and comprehensive_suggestions.get('optimization_suggestions'):
                    formatted = self.intelligent_optimizer.format_suggestions_for_report(comprehensive_suggestions)
                    if formatted and formatted.strip() and formatted != "暂无优化建议":
                        return formatted
            except Exception:
                # 智能模块不可用时，继续使用兜底逻辑
                pass
        
        # 兜底的通用建议
        fallback = [
            "1. 智能诊断: 暂无足够信息生成具体诊断结果，但SQL仍存在优化空间",
            "2. 智能优化建议:",
            "• 使用EXPLAIN分析执行计划，确认是否存在全表扫描",
            "• 确认WHERE/JOIN字段均已建立合适索引",
            "• 避免SELECT *，只返回必要字段",
            "3. 预期效果: 预计平均查询时间可降低50%以上"
        ]
        return "\n".join(fallback)
    
    def check_indexes_exist(self, database: str, table_name: str, where_fields: list, join_fields: list, order_by_fields: list, query: Optional[dict] = None) -> bool:
        """检查所有相关字段是否都有索引（占位方法）"""
        # TODO: 从主文件复制完整实现
        return False
    
    def check_composite_index_exists(self, existing_indexed_fields: set, composite_fields: list) -> bool:
        """检查是否已有复合索引覆盖指定的字段组合（占位方法）"""
        # TODO: 从主文件复制完整实现
        return False
    
